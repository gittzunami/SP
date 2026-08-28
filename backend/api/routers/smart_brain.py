"""api/routers/smart_brain.py — Smart Brain (AI analysis) endpoints."""

from __future__ import annotations

import io
import logging
from datetime import datetime, timedelta, timezone
from typing import Any, Dict, List

from fastapi import APIRouter, Body, Depends, File, HTTPException, Query, UploadFile
from fastapi import Path as FPath
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from database import get_db

logger = logging.getLogger("smart_brain")
router = APIRouter(prefix="/api/smart-brain", tags=["SmartBrain"])
webhook_router = APIRouter(prefix="/api/webhook/smart-brain", tags=["SmartBrainWebhook"])


# ── Request schemas ───────────────────────────────────────────────────────────

class SmartBrainRunRequest(BaseModel):
    prompt:      str
    run_ids:     List[int]
    max_per_run: int | None = None
    keyword:     str        = ""


class SmartBrainBatchRunRequest(BaseModel):
    batch_id:   str | None = None
    prompt:     str | None = None
    chunk_size: int | None = 500


class SmartBrainEnhanceSingleRequest(BaseModel):
    prompt: str
    record: Dict[str, Any]


class SmartBrainDirectRunRequest(BaseModel):
    prompt:  str
    records: List[Dict[str, Any]]


class SmartBrainExportRequest(BaseModel):
    text:  str
    title: str = "Smart Brain Analysis"


# ── LLM field stripping (remove high-cardinality noise before analysis) ───────

_LLM_STRIP_FIELDS: dict[str, set] = {
    "reddit":        {"score"},
    "tiktok":        {"plays", "likes"},
    "autodesk":      {"kudos"},
    "stackexchange": {"score"},
    "instagram":     {"likes", "comments_count"},
    "twitter":       {"retweets", "replies", "avatar"},
    "quora":         {"answer_count", "body"},
    "facebook":      {"likes_count", "comments_count"},
}


def _smart_brain_records_for_runs(db, runs, max_per_run: int | None = None) -> list[dict]:
    from services.search_service import (
        _reddit_post_preview, _tiktok_post_preview, _edugeek_post_preview,
        _autodesk_post_preview, _se_question_preview, _gnews_preview,
        _instagram_post_preview, _spiceworks_preview, _twitter_tweet_preview,
        _quora_question_preview, _facebook_post_preview,
    )
    from db_models import (
        RedditPost, TikTokPost, EduGeekPost, AutodeskPost, StackExchangeQuestion,
        GoogleNewsArticle, InstagramPost, SpiceworksPost, TwitterTweet, QuoraQuestion,
        FacebookPost,
    )
    SOURCE_MAP = {
        "reddit":        (RedditPost,            _reddit_post_preview),
        "tiktok":        (TikTokPost,            _tiktok_post_preview),
        "edugeek":       (EduGeekPost,           _edugeek_post_preview),
        "autodesk":      (AutodeskPost,          _autodesk_post_preview),
        "stackexchange": (StackExchangeQuestion, _se_question_preview),
        "instagram":     (InstagramPost,         _instagram_post_preview),
        "spiceworks":    (SpiceworksPost,        _spiceworks_preview),
        "twitter":       (TwitterTweet,          _twitter_tweet_preview),
        "quora":         (QuoraQuestion,         _quora_question_preview),
        "facebook":      (FacebookPost,          _facebook_post_preview),
    }
    rows = []
    for run in runs:
        scraper = (run.scraper or "").lower()
        if scraper == "google_news":
            q = db.query(GoogleNewsArticle)
            if run.keyword:
                q = q.filter(GoogleNewsArticle.search_query == run.keyword)
            elif run.scraped_at:
                window_start = run.scraped_at - timedelta(minutes=10)
                window_end   = run.scraped_at + timedelta(minutes=10)
                q = q.filter(
                    GoogleNewsArticle.scraped_at >= window_start,
                    GoogleNewsArticle.scraped_at <= window_end,
                )
            if max_per_run is not None:
                q = q.limit(max_per_run)
            recs = q.all()
            rows.extend(_gnews_preview(r, db) for r in recs)
            continue
        entry = SOURCE_MAP.get(scraper)
        if not entry:
            continue
        model, preview_fn = entry
        try:
            q = db.query(model).filter(model.run_id == run.id)
            if max_per_run is not None:
                q = q.limit(max_per_run)
            recs  = q.all()
            strip = _LLM_STRIP_FIELDS.get(scraper, set())
            for r in recs:
                row = preview_fn(r, db)
                for f in strip:
                    row.pop(f, None)
                rows.append(row)
        except Exception as exc:
            logger.warning("smart_brain: failed to fetch %s run %s: %s", scraper, run.id, exc)
    return rows


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.get("/sessions")
def smart_brain_sessions(
    from_date: str | None = Query(None),
    limit:     int        = Query(10, le=500),
    order:     str        = Query("desc"),
    db:        Session    = Depends(get_db),
):
    from db_models import ScrapeRun
    asc_order = order == "asc"
    q = db.query(ScrapeRun)
    if from_date:
        try:
            dt = datetime.fromisoformat(from_date)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            q = q.filter(ScrapeRun.scraped_at >= dt)
        except Exception:
            pass
    q = q.order_by(ScrapeRun.scraped_at.asc() if asc_order else ScrapeRun.scraped_at.desc())
    runs = q.all()

    sessions: list[list] = []
    current:  list       = []
    prev_dt              = None
    for run in runs:
        run_dt = run.scraped_at
        if run_dt and run_dt.tzinfo is None:
            run_dt = run_dt.replace(tzinfo=timezone.utc)
        if prev_dt is None or abs((prev_dt - (run_dt or prev_dt)).total_seconds()) <= 600:
            current.append(run)
            prev_dt = run_dt or prev_dt
        else:
            sessions.append(current)
            current = [run]
            prev_dt = run_dt
    if current:
        sessions.append(current)

    sessions = sessions[:limit]
    result = []
    for i, grp in enumerate(sessions):
        first    = grp[0]
        first_dt = first.scraped_at
        result.append({
            "session_index": i,
            "scraped_at":    first_dt.isoformat() if first_dt else None,
            "scraper_count": len(grp),
            "total_items":   sum(r.total_items or 0 for r in grp),
            "runs": [
                {
                    "run_id":      r.id,
                    "scraper":     r.scraper,
                    "keyword":     r.keyword or "",
                    "scraped_at":  r.scraped_at.isoformat() if r.scraped_at else None,
                    "total_items": r.total_items or 0,
                }
                for r in grp
            ],
        })
    return {"sessions": result}


@router.post("/parse-file")
async def smart_brain_parse_file(file: UploadFile = File(...)):
    content  = await file.read()
    filename = file.filename or ""
    ext      = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    text     = ""
    try:
        if ext in ("txt", "mx"):
            text = content.decode("utf-8", errors="replace")
        elif ext == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text   = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in ("doc", "docx"):
            from docx import Document
            doc  = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            raise HTTPException(400, "Unsupported file type. Upload a .txt, .pdf, .docx, or .mx file.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, f"Failed to parse file: {exc}")
    return {"text": text.strip(), "filename": filename}


@router.post("/run")
def smart_brain_run(body: SmartBrainRunRequest, db: Session = Depends(get_db)):
    if not body.prompt.strip():
        raise HTTPException(400, "Prompt must not be empty.")
    if not body.run_ids:
        raise HTTPException(400, "Select at least one scraper run.")

    from db_models import ScrapeRun
    runs = db.query(ScrapeRun).filter(ScrapeRun.id.in_(body.run_ids)).all()
    if not runs:
        raise HTTPException(404, "None of the selected run IDs were found.")

    run_limit = max(1, body.max_per_run) if body.max_per_run is not None else None
    data_rows = _smart_brain_records_for_runs(db, runs, max_per_run=run_limit)
    if not data_rows:
        raise HTTPException(400, "No records found for the selected runs.")

    from llm_service import feed_to_llm, enhance_prompt
    try:
        enhanced = enhance_prompt(db, body.prompt, sample_rows=data_rows[:3])
        ep       = enhanced.get("enhanced_prompt", body.prompt)
    except Exception:
        ep = body.prompt

    try:
        result = feed_to_llm(db, ep, data_rows, keyword=body.keyword or body.prompt[:80])
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))

    return {
        "response":     result["response"],
        "provider":     result["provider"],
        "model":        result["model"],
        "tokens_used":  result.get("tokens_used", 0),
        "cost_usd":     result.get("cost_usd", 0.0),
        "records_used": len(data_rows),
    }


@router.post("/export-docx")
def smart_brain_export_docx(body: SmartBrainExportRequest):
    from docx import Document

    doc = Document()
    doc.add_heading(body.title, level=0)

    for line in body.text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 40)
        else:
            p     = doc.add_paragraph()
            parts = stripped.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="smart_brain_analysis.docx"'},
    )


@router.post("/enhance-single")
def smart_brain_enhance_single(body: SmartBrainEnhanceSingleRequest, db: Session = Depends(get_db)):
    if not body.prompt.strip():
        raise HTTPException(400, "Prompt must not be empty.")
    from llm_service import enhance_prompt
    try:
        result = enhance_prompt(db, body.prompt, sample_rows=[body.record])
        return {
            "enhanced_prompt":  result.get("enhanced_prompt", body.prompt),
            "summary_for_user": result.get("summary_for_user", ""),
        }
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))


@router.post("/run-direct")
def smart_brain_run_direct(body: SmartBrainDirectRunRequest, db: Session = Depends(get_db)):
    if not body.prompt.strip():
        raise HTTPException(400, "Prompt must not be empty.")
    if not body.records:
        raise HTTPException(400, "No records provided.")
    cleaned = []
    for rec in body.records:
        source = rec.get("source", "")
        strip  = _LLM_STRIP_FIELDS.get(source, set())
        cleaned.append({k: v for k, v in rec.items() if k not in strip})
    from llm_service import feed_to_llm
    try:
        result = feed_to_llm(db, body.prompt, cleaned)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    return {
        "response":     result["response"],
        "provider":     result["provider"],
        "model":        result["model"],
        "tokens_used":  result.get("tokens_used", 0),
        "cost_usd":     result.get("cost_usd", 0.0),
        "records_used": len(cleaned),
    }


@router.post("/history")
def save_smart_brain_analysis(body: dict = Body(...), db: Session = Depends(get_db)):
    from db_models import SmartBrainAnalysis
    entry = SmartBrainAnalysis(
        result          = body.get("result", ""),
        provider        = body.get("provider", ""),
        model           = body.get("model", ""),
        tokens_used     = body.get("tokens_used", 0),
        cost_usd        = body.get("cost_usd", 0.0),
        enhanced_prompt = body.get("enhanced_prompt", ""),
        prompt_used     = body.get("prompt_used", ""),
        record_count    = body.get("record_count", 0),
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return {"id": entry.id, "created_at": entry.created_at.isoformat()}


@router.get("/history")
def get_smart_brain_history(limit: int = Query(100, le=200), db: Session = Depends(get_db)):
    from db_models import SmartBrainAnalysis
    rows = db.query(SmartBrainAnalysis).order_by(SmartBrainAnalysis.created_at.desc()).limit(limit).all()
    return {"analyses": [
        {
            "id":              r.id,
            "result":          r.result,
            "provider":        r.provider,
            "model":           r.model,
            "tokens_used":     r.tokens_used,
            "cost_usd":        r.cost_usd,
            "enhanced_prompt": r.enhanced_prompt,
            "prompt_used":     r.prompt_used,
            "record_count":    r.record_count,
            "timestamp":       r.created_at.isoformat() if r.created_at else "",
        }
        for r in rows
    ]}


@router.delete("/history/{entry_id}")
def delete_smart_brain_analysis(entry_id: int = FPath(...), db: Session = Depends(get_db)):
    from db_models import SmartBrainAnalysis
    entry = db.query(SmartBrainAnalysis).filter_by(id=entry_id).first()
    if not entry:
        raise HTTPException(404, "Analysis not found.")
    db.delete(entry)
    db.commit()
    return {"deleted": entry_id}


# ══════════════════════════════════════════════════════════════════════════════
#  Saved Prompts — /api/smart-brain/prompts
#  Each "Save Prompt" on the Smart Brain page creates a new row.
#  The popup and Smart Brain page both read from this table.
# ══════════════════════════════════════════════════════════════════════════════

class SavedPromptSchema(BaseModel):
    text:  str
    label: str = ""


@router.get("/prompts")
def list_saved_prompts(
    limit: int = Query(100, le=500),
    db:    Session = Depends(get_db),
):
    from db_models import SavedPrompt, Base
    try:
        rows = (
            db.query(SavedPrompt)
            .order_by(SavedPrompt.created_at.desc())
            .limit(limit)
            .all()
        )
    except Exception:
        db.rollback()
        Base.metadata.create_all(bind=db.get_bind(), tables=[SavedPrompt.__table__])
        rows = (
            db.query(SavedPrompt)
            .order_by(SavedPrompt.created_at.desc())
            .limit(limit)
            .all()
        )
    return {
        "prompts": [
            {
                "id":         r.id,
                "text":       r.text,
                "label":      r.label or "",
                "created_at": r.created_at.isoformat() if r.created_at else "",
                "updated_at": r.updated_at.isoformat() if r.updated_at else "",
            }
            for r in rows
        ]
    }


@router.post("/prompts")
def create_saved_prompt(
    body: SavedPromptSchema,
    db:   Session = Depends(get_db),
):
    from db_models import SavedPrompt, Base
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc)
    prompt = SavedPrompt(
        text       = body.text.strip(),
        label      = body.label.strip() if body.label else None,
        created_at = now,
        updated_at = now,
    )
    try:
        db.add(prompt)
        db.commit()
    except Exception:
        db.rollback()
        Base.metadata.create_all(bind=db.get_bind(), tables=[SavedPrompt.__table__])
        db.add(prompt)
        db.commit()

    db.refresh(prompt)
    return {
        "id":         prompt.id,
        "text":       prompt.text,
        "label":      prompt.label or "",
        "created_at": prompt.created_at.isoformat() if prompt.created_at else now.isoformat(),
        "updated_at": prompt.updated_at.isoformat() if prompt.updated_at else now.isoformat(),
    }


@router.put("/prompts/{prompt_id}")
def update_saved_prompt(
    prompt_id: int = FPath(...),
    body: SavedPromptSchema = Body(...),
    db:   Session = Depends(get_db),
):
    from db_models import SavedPrompt, Base
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc)
    try:
        prompt = db.query(SavedPrompt).filter_by(id=prompt_id).first()
    except Exception:
        db.rollback()
        Base.metadata.create_all(bind=db.get_bind(), tables=[SavedPrompt.__table__])
        prompt = db.query(SavedPrompt).filter_by(id=prompt_id).first()

    if not prompt:
        raise HTTPException(404, "Prompt not found.")
    prompt.text       = body.text.strip()
    prompt.label      = body.label.strip() if body.label else None
    prompt.updated_at = now
    db.commit()
    db.refresh(prompt)
    return {
        "id":         prompt.id,
        "text":       prompt.text,
        "label":      prompt.label or "",
        "created_at": prompt.created_at.isoformat() if prompt.created_at else now.isoformat(),
        "updated_at": prompt.updated_at.isoformat() if prompt.updated_at else now.isoformat(),
    }


@router.delete("/prompts/{prompt_id}")
def delete_saved_prompt(
    prompt_id: int = FPath(...),
    db:        Session = Depends(get_db),
):
    from db_models import SavedPrompt
    prompt = db.query(SavedPrompt).filter_by(id=prompt_id).first()
    if not prompt:
        raise HTTPException(404, "Prompt not found.")
    db.delete(prompt)
    db.commit()
    return {"deleted": prompt_id}


# ══════════════════════════════════════════════════════════════════════════════
#  Batch Inspection & Manual Map-Reduce Endpoints (Zero-Auth & Authenticated)
# ══════════════════════════════════════════════════════════════════════════════

def _list_batches_impl(db: Session) -> dict:
    from db_models import ScrapeRun

    runs = db.query(ScrapeRun).filter(ScrapeRun.batch_id.isnot(None)).order_by(ScrapeRun.scraped_at.desc()).all()
    if not runs:
        return {"total_batches": 0, "batches": []}

    batches_map: dict[str, dict] = {}
    for r in runs:
        bid = r.batch_id
        if bid not in batches_map:
            batches_map[bid] = {
                "batch_id":      bid,
                "total_records": 0,
                "run_count":     0,
                "scrapers":      set(),
                "keywords":      set(),
                "started_at":    r.scraped_at.isoformat() if r.scraped_at else None,
                "run_ids":       [],
            }
        b = batches_map[bid]
        b["total_records"] += (r.total_items or 0)
        b["run_count"]     += 1
        b["run_ids"].append(r.id)
        if r.scraper:
            b["scrapers"].add(r.scraper)
        if r.keyword:
            b["keywords"].add(r.keyword)

    result = []
    for bid, data in batches_map.items():
        result.append({
            "batch_id":      bid,
            "total_records": data["total_records"],
            "run_count":     data["run_count"],
            "scrapers":      sorted(list(data["scrapers"])),
            "keywords":      sorted(list(data["keywords"])),
            "started_at":    data["started_at"],
            "run_ids":       data["run_ids"],
        })

    return {"total_batches": len(result), "batches": result}


def _execute_batch_smart_brain(db: Session, batch_id: str | None = None, custom_prompt: str | None = None, chunk_size: int = 500) -> dict:
    from db_models import ScrapeRun, SavedPrompt, SmartBrainAnalysis
    from llm_service import feed_to_llm

    # If batch_id is omitted or "latest", auto-detect the most recent scrape batch in DB
    if not batch_id or str(batch_id).strip().lower() in ("", "latest", "none", "null"):
        latest_run = db.query(ScrapeRun).filter(ScrapeRun.batch_id.isnot(None)).order_by(ScrapeRun.scraped_at.desc()).first()
        if not latest_run:
            raise HTTPException(status_code=404, detail="No scrape batches found in database.")
        batch_id = latest_run.batch_id
        logger.info("Auto-detected latest batch_id: %s", batch_id)

    # 1. Fetch all runs matching this batch
    runs = db.query(ScrapeRun).filter(ScrapeRun.batch_id == batch_id).all()
    if not runs:
        raise HTTPException(status_code=404, detail=f"No scrape runs found for batch_id: {batch_id}")

    # 2. Extract records across all runs in this batch (unlimited to include all records)
    data_rows = _smart_brain_records_for_runs(db, runs, max_per_run=None)
    if not data_rows:
        raise HTTPException(status_code=400, detail=f"No scraped records found in database for batch_id: {batch_id}")

    # 3. Resolve Prompt
    prompt_text = (custom_prompt or "").strip()
    if not prompt_text:
        saved = db.query(SavedPrompt).order_by(SavedPrompt.created_at.desc()).first()
        if saved and saved.text:
            prompt_text = saved.text
        else:
            prompt_text = "Perform a thorough intelligence analysis. Identify key trends, sentiment distribution, critical findings, and actionable recommendations."

    logger.info(
        "Executing batch Smart Brain on batch %s (%d records, chunk_size=%d)",
        batch_id[:8], len(data_rows), chunk_size,
    )

    # 4. Run Map-Reduce Analysis (handles 500 chunks + 250 recursive fallback)
    try:
        result = feed_to_llm(
            db=db,
            enhanced_prompt=prompt_text,
            data_rows=data_rows,
            keyword=f"Batch {batch_id[:8]}",
        )
    except RuntimeError as exc:
        logger.error("Batch Smart Brain failed for batch %s: %s", batch_id[:8], exc)
        raise HTTPException(status_code=500, detail=str(exc))

    # 5. Persist to PostgreSQL smart_brain_analyses
    entry = SmartBrainAnalysis(
        result          = result.get("response", ""),
        provider        = result.get("provider", "openai"),
        model           = result.get("model", ""),
        tokens_used     = result.get("tokens_used", 0),
        cost_usd        = result.get("cost_usd", 0.0),
        enhanced_prompt = "",
        prompt_used     = prompt_text,
        record_count    = len(data_rows),
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)

    logger.info(
        "Batch Smart Brain saved (id=%d, records=%d, tokens=%d, cost=$%.6f)",
        entry.id, len(data_rows), entry.tokens_used, entry.cost_usd,
    )

    return {
        "status":           "completed",
        "analysis_id":      entry.id,
        "batch_id":         batch_id,
        "record_count":     len(data_rows),
        "tokens_used":      entry.tokens_used,
        "cost_usd":         entry.cost_usd,
        "provider":         entry.provider,
        "model":            entry.model,
        "prompt_used":      prompt_text,
        "response":         entry.result,
        "created_at":       entry.created_at.isoformat() if entry.created_at else None,
    }


# ── Authenticated Routes ───────────────────────────────────────────────────────

@router.get("/batches")
def list_batches_auth(db: Session = Depends(get_db)):
    return _list_batches_impl(db)


@router.post("/run-batch")
def run_batch_auth(payload: SmartBrainBatchRunRequest = Body(default_factory=SmartBrainBatchRunRequest), db: Session = Depends(get_db)):
    return _execute_batch_smart_brain(db, payload.batch_id, payload.prompt, payload.chunk_size or 500)


@router.post("/run-latest")
def run_latest_auth(db: Session = Depends(get_db)):
    return _execute_batch_smart_brain(db, batch_id=None)


@router.post("/batches/{batch_id}/run")
def run_batch_path_auth(batch_id: str = FPath(...), db: Session = Depends(get_db)):
    return _execute_batch_smart_brain(db, batch_id)


# ── Webhook / Zero-Auth Routes (No Authorization header needed) ────────────────

@webhook_router.get("/batches")
def list_batches_webhook(db: Session = Depends(get_db)):
    return _list_batches_impl(db)


@webhook_router.post("/run-batch")
def run_batch_webhook(payload: SmartBrainBatchRunRequest = Body(default_factory=SmartBrainBatchRunRequest), db: Session = Depends(get_db)):
    return _execute_batch_smart_brain(db, payload.batch_id, payload.prompt, payload.chunk_size or 500)


@webhook_router.post("/run-latest")
def run_latest_webhook(db: Session = Depends(get_db)):
    return _execute_batch_smart_brain(db, batch_id=None)


@webhook_router.post("/batches/{batch_id}/run")
def run_batch_path_webhook(batch_id: str = FPath(...), db: Session = Depends(get_db)):
    return _execute_batch_smart_brain(db, batch_id)
