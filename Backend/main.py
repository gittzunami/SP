"""
main.py — TrendSense Unified Scraper API v4.2
=============================================

Start:  uvicorn main:app --reload --port 8000
Docs:   http://localhost:8000/docs
"""

from __future__ import annotations

import os
from dotenv import load_dotenv
load_dotenv()

import json
import logging
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List

from datetime import timedelta

from fastapi import BackgroundTasks, Body, Depends, FastAPI, File, HTTPException, Query, Request, UploadFile
from fastapi import Path as FPath
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import OAuth2PasswordBearer
from pydantic import BaseModel
from sqlalchemy.orm import Session
from jose import JWTError, jwt

import scheduler_service as sched
import database
from database import get_db
from models import (
    AutodeskConfig, EduGeekConfig, RedditConfig,
    StackExchangeConfig, TikTokConfig,
    TwitterConfig, InstagramConfig, GoogleNewsConfig,
    SpiceworksConfig, QuoraConfig, FacebookConfig,
    RunRequest, RunResponse, ScheduleRequest,
)

logging.basicConfig(
    level   = logging.INFO,
    format  = "%(asctime)s  %(levelname)-8s  %(name)s — %(message)s",
    datefmt = "%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("api")

RESULTS_DIR = Path("results")
RESULTS_DIR.mkdir(exist_ok=True)

# ── Auth config ────────────────────────────────────────────────────────────────
JWT_SECRET    = os.environ.get("JWT_SECRET_KEY", "")
if not JWT_SECRET:
    JWT_SECRET = "trendsense-dev-secret-change-in-prod"
    logger.warning(
        "⚠️  JWT_SECRET_KEY not set in .env — using insecure default. "
        "Set JWT_SECRET_KEY=<random-64-char-string> in your .env file before going to production."
    )
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_H  = int(os.environ.get("JWT_EXPIRE_HOURS", "8"))

LOGIN_USERNAME = os.environ.get("LOGIN_USERNAME", "").strip()
LOGIN_PASSWORD = os.environ.get("LOGIN_PASSWORD", "").strip()
if not LOGIN_USERNAME or not LOGIN_PASSWORD:
    logger.warning(
        "⚠️  LOGIN_USERNAME or LOGIN_PASSWORD not set in .env — "
        "add LOGIN_USERNAME=youruser and LOGIN_PASSWORD=yourpass to enable login."
    )

# Paths/prefixes that bypass JWT check
_AUTH_SKIP_EXACT   = {"/api/auth/login", "/docs", "/openapi.json", "/redoc", "/", "/api/health"}
_AUTH_SKIP_PREFIX  = ("/webhook/", "/api/webhook/")   # Power Automate callbacks


class _LoginRequest(BaseModel):
    username: str
    password: str

class SmartBrainRunRequest(BaseModel):
    prompt:      str
    run_ids:     List[int]
    max_per_run: int  = 100
    keyword:     str  = ""

class SmartBrainEnhanceSingleRequest(BaseModel):
    prompt: str
    record: Dict[str, Any]

class SmartBrainDirectRunRequest(BaseModel):
    prompt:  str
    records: List[Dict[str, Any]]

class SmartBrainExportRequest(BaseModel):
    text:  str
    title: str = "Smart Brain Analysis"

VALID_SCRAPERS = {
    "reddit", "tiktok", "edugeek", "stackexchange", "autodesk",
    "twitter", "instagram", "google_news", "spiceworks", "quora", "facebook",
}

task_registry:  Dict[str, Dict[str, Any]] = {}
scraper_status: Dict[str, Dict[str, Any]] = {
    s: {"last_run": None, "last_file": None, "total_runs": 0, "running": False,
        "last_total_items": None, "last_newsletters_created": None}
    for s in VALID_SCRAPERS
}


def _get_scraper(name: str):
    if name == "reddit":        from scrapers import reddit;              return reddit
    if name == "tiktok":        from scrapers import tiktok;              return tiktok
    if name == "edugeek":       from scrapers import edugeek;             return edugeek
    if name == "stackexchange": from scrapers import stackexchange;       return stackexchange
    if name == "autodesk":      from scrapers import autodesk;            return autodesk
    if name == "twitter":       from scrapers import getxapi_twitter;     return getxapi_twitter
    if name == "instagram":     from scrapers import apify_instagram;     return apify_instagram
    if name == "google_news":   from scrapers import scrappa_google_news; return scrappa_google_news
    if name == "spiceworks":    from scrapers import spiceworks;          return spiceworks
    if name == "quora":         from scrapers import quora;               return quora
    if name == "facebook":      from scrapers import facebook;             return facebook
    raise ValueError(f"Unknown scraper: {name!r}")


# ══════════════════════════════════════════════════════════════════════════════
#  Lifespan — restores scraper last_run from DB on startup
# ══════════════════════════════════════════════════════════════════════════════

@asynccontextmanager
async def lifespan(app: FastAPI):
    database.init_db()
    sched.get_scheduler()

    # Restore scraper last_run from DB so cards show correct state after restart
    if database.SessionLocal is not None:
        try:
            db = database.SessionLocal()
            from db_models import ScrapeRun
            from sqlalchemy import func as _func

            # Subquery: most recent scraped_at per scraper
            subq = (
                db.query(
                    ScrapeRun.scraper,
                    _func.max(ScrapeRun.scraped_at).label("max_at"),
                )
                .group_by(ScrapeRun.scraper)
                .subquery()
            )
            # Join to get the full row (including total_items) for that most-recent run
            rows = (
                db.query(ScrapeRun)
                .join(subq,
                      (ScrapeRun.scraper == subq.c.scraper) &
                      (ScrapeRun.scraped_at == subq.c.max_at))
                .all()
            )
            for row in rows:
                if row.scraper in scraper_status and row.scraped_at:
                    scraper_status[row.scraper]["last_run"]         = row.scraped_at.isoformat()
                    scraper_status[row.scraper]["last_total_items"] = row.total_items or 0
            db.close()
            logger.info("Restored scraper_status from DB (%d scrapers)", len(rows))
        except Exception as exc:
            logger.warning("Could not restore scraper_status from DB: %s", exc)

    # One-time cleanup: normalize all existing keywords to lowercase
    if database.SessionLocal is not None:
        try:
            db = database.SessionLocal()
            from db_models import ScrapeRun
            from sqlalchemy import update, func as _sfunc
            db.execute(
                update(ScrapeRun)
                .where(ScrapeRun.keyword.isnot(None), ScrapeRun.keyword != "")
                .values(keyword=_sfunc.lower(ScrapeRun.keyword))
            )
            db.commit()
            db.close()
        except Exception as exc:
            logger.warning("Keyword lowercase cleanup failed: %s", exc)

    yield
    sched.shutdown()


# ══════════════════════════════════════════════════════════════════════════════
#  App
# ══════════════════════════════════════════════════════════════════════════════

app = FastAPI(
    title    = "TrendSense Scraper API",
    version  = "4.2.0",
    lifespan = lifespan,
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], allow_credentials=True,
    allow_methods=["*"], allow_headers=["*"],
)

from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse as _JSONResponse

@app.exception_handler(RequestValidationError)
async def _validation_error_handler(request: Request, exc: RequestValidationError):
    logger.error(
        "422 Validation error on %s %s — %s",
        request.method, request.url.path, exc.errors(),
    )
    return _JSONResponse(status_code=422, content={"detail": exc.errors()})


@app.middleware("http")
async def _jwt_middleware(request: Request, call_next):
    path = request.url.path

    # Always allow: CORS preflight, public paths, webhook callbacks
    if (request.method == "OPTIONS"
            or path in _AUTH_SKIP_EXACT
            or any(path.startswith(p) for p in _AUTH_SKIP_PREFIX)):
        return await call_next(request)

    # Guard EVERYTHING else — including /search, /export, /export/selected
    raw   = request.headers.get("Authorization", "")
    token = raw[7:] if raw.startswith("Bearer ") else ""
    if not token:
        return JSONResponse(status_code=401, content={"detail": "Not authenticated"})
    try:
        jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except JWTError:
        return JSONResponse(status_code=401, content={"detail": "Invalid or expired token"})

    return await call_next(request)


# ── Auth endpoints ────────────────────────────────────────────────────────────

@app.post("/api/auth/login", tags=["Auth"])
def login(payload: _LoginRequest):
    if not LOGIN_USERNAME or not LOGIN_PASSWORD:
        raise HTTPException(
            status_code=503,
            detail="Login not configured — set LOGIN_USERNAME and LOGIN_PASSWORD in .env",
        )
    if payload.username != LOGIN_USERNAME or payload.password != LOGIN_PASSWORD:
        raise HTTPException(status_code=401, detail="Invalid username or password")

    token = jwt.encode(
        {"sub": LOGIN_USERNAME, "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRE_H)},
        JWT_SECRET,
        algorithm=JWT_ALGORITHM,
    )
    return {"access_token": token, "token_type": "bearer", "username": LOGIN_USERNAME}


@app.get("/api/auth/me", tags=["Auth"])
def me(request: Request):
    raw   = request.headers.get("Authorization", "")
    token = raw[7:] if raw.startswith("Bearer ") else ""
    try:
        data = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid token")
    return {"username": data.get("sub")}


# ══════════════════════════════════════════════════════════════════════════════
#  Background runner
# ══════════════════════════════════════════════════════════════════════════════

def _filter_by_date(scraper: str, result: dict, since_date: str) -> dict:
    """Remove items whose date is before since_date. Applied after scraping and before DB save."""
    if not since_date:
        return result
    try:
        since = datetime.fromisoformat(since_date).replace(tzinfo=timezone.utc)
    except ValueError:
        logger.warning("Invalid since_date %r — skipping date filter", since_date)
        return result

    def _keep(item: dict, *date_keys: str) -> bool:
        for key in date_keys:
            raw = item.get(key)
            if not raw:
                continue
            try:
                raw_s = str(raw)
                # Handle Twitter-style date: "Thu Mar 26 13:04:38 +0000 2026"
                if raw_s.count(":") == 2 and "+" in raw_s and len(raw_s) > 25:
                    dt = datetime.strptime(raw_s, "%a %b %d %H:%M:%S %z %Y")
                else:
                    dt = datetime.fromisoformat(raw_s.replace("Z", "+00:00"))
                if dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
                return dt >= since
            except Exception:
                continue
        return True  # no parseable date → keep

    result = dict(result)  # shallow copy so we don't mutate the original

    if scraper in ("reddit", "autodesk"):
        posts = [p for p in result.get("posts", []) if _keep(p, "created_at")]
        result["posts"] = posts
        result["total_posts"] = len(posts)

    elif scraper == "edugeek":
        cats = result.get("categories", {})
        filtered = {cat: [i for i in items if _keep(i, "created_at")] for cat, items in cats.items()}
        result["categories"]      = filtered
        result["total_items"]     = sum(len(v) for v in filtered.values())
        result["category_counts"] = {k: len(v) for k, v in filtered.items()}

    elif scraper == "stackexchange":
        qs = [q for q in result.get("questions", []) if _keep(q, "created_at")]
        result["questions"]      = qs
        result["total_questions"] = len(qs)

    elif scraper == "twitter":
        tweets = [t for t in result.get("tweets", [])
                  if _keep(t, "created_at", "date", "createdAt")]
        result["tweets"]       = tweets
        result["total_tweets"] = len(tweets)

    elif scraper == "google_news":
        articles = [a for a in result.get("articles", [])
                    if _keep(a, "publishedAt", "published_at", "date", "datePublished")]
        result["articles"]       = articles
        result["total_articles"] = len(articles)

    elif scraper == "spiceworks":
        posts = [p for p in result.get("posts", []) if _keep(p, "date")]
        result["posts"]       = posts
        result["total_items"] = len(posts)

    logger.info("Date filter (%s, since=%s): result after filter applied", scraper, since_date)
    return result


def _run_scraper(task_id: str, scraper: str, cfg) -> None:
    task_registry[task_id]["status"]   = "running"
    scraper_status[scraper]["running"] = True

    spend_db = database.SessionLocal() if database.SessionLocal is not None else None

    # ── Per-scraper budget block check ────────────────────────────────────────
    if spend_db:
        try:
            from services.spending_service import get_scraper_budget_status
            bs = get_scraper_budget_status(spend_db).get(scraper, {})
            if bs.get("is_blocked"):
                reason = (
                    "No budget allocated — set a budget in Cost Governance"
                    if bs.get("no_budget") else
                    f"Budget limit reached ({bs.get('pct', 0):.1f}% of ${bs.get('budget_usd', 0):.2f})"
                )
                err_msg = f"Scraper blocked: {reason}"
                logger.warning("Task %s (%s) blocked before start — %s", task_id[:8], scraper, reason)
                task_registry[task_id].update({"status": "failed", "error": err_msg,
                                               "finished_at": datetime.now(tz=timezone.utc).isoformat()})
                scraper_status[scraper]["running"] = False
                try:
                    from db_models import TaskHistory
                    row = spend_db.query(TaskHistory).filter_by(task_id=task_id).first()
                    if row:
                        row.status = "failed"; row.error = err_msg[:500]
                        row.finished_at = datetime.now(tz=timezone.utc)
                        spend_db.commit()
                except Exception:
                    pass
                spend_db.close()
                return
        except Exception as exc:
            logger.warning("Budget block check failed for %s: %s", scraper, exc)

    try:
        mod = _get_scraper(scraper)

        # ── Call the right entry-point ─────────────────────────────────────────
        if scraper == "twitter":
            result = mod.run_twitter(
                keywords=cfg.keywords, max_tweets=cfg.max_tweets,
                lang=getattr(cfg, "lang", "en"),
                task_id=task_id,
            )
        elif scraper == "instagram":
            result = mod.run_instagram(
                keywords=cfg.keywords, results_limit=cfg.results_limit,
                task_id=task_id,
            )
        elif scraper == "google_news":
            result = mod.run_google_news(
                keywords    = cfg.keywords,
                max_results = cfg.max_results,
                task_id     = task_id,
                db          = spend_db,
            )
        else:
            result = mod.run(cfg)

        # ── Apply date filter (first pass, before DB save) ─────────────────────
        since_date = getattr(cfg, "since_date", None)
        if since_date:
            result = _filter_by_date(scraper, result, since_date)

        # ── Persist to PostgreSQL (skip google_news — saved after webhook approval) ──
        # Capture scraped counts BEFORE the DB save overwrites them — cost is
        # charged on items fetched from the source, not DB-deduplicated rows.
        _scraped_counts = {
            k: result[k]
            for k in ("total_posts", "total_questions", "total_tweets", "total_articles", "total_items")
            if k in result
        }
        if spend_db is not None and scraper != "google_news":
            from services.db_writer import save
            try:
                actual_saved = save(scraper, spend_db, result, task_id, since_date=since_date)
                for _key in ("total_posts", "total_questions", "total_tweets", "total_articles", "total_items"):
                    if _key in result:
                        result[_key] = actual_saved
                        break
                logger.info("DB write complete for task %s (%s): %d saved (scraped: %s)",
                            task_id[:8], scraper, actual_saved, _scraped_counts)
            except Exception as exc:
                logger.error("DB write error for %s: %s", scraper, exc)
                spend_db.rollback()

        # ── Record spending ────────────────────────────────────────────────────
        # Apify scrapers (google_news, twitter, instagram) record cost HERE
        # (not inside their scraper functions) so the env-var rate is always
        # used and cost is always captured in a clean, independent transaction.
        if spend_db is not None:
            try:
                from services.spending_service import (
                    record_apify_spend,
                    record_reddit_spend, record_autodesk_spend,
                    record_scrapecreators_spend, record_scrapingbee_spend,
                )
                keyword = getattr(cfg, "keyword", "") or ""
                if not keyword and hasattr(cfg, "keywords"):
                    keyword = (cfg.keywords or [""])[0]

                if scraper == "google_news":
                    from services.spending_service import record_scrappa_spend
                    scrappa_stats = result.get("_scrappa_run_stats") or {}
                    record_scrappa_spend(
                        db             = spend_db,
                        requests_made  = scrappa_stats.get("requests_made", 0),
                        articles_found = scrappa_stats.get("articles_found", result.get("total_articles", 0)),
                        keyword        = keyword,
                        task_id        = task_id,
                    )
                elif scraper == "twitter":
                    from services.spending_service import record_getxapi_spend
                    gx_stats = result.get("_getxapi_run_stats") or {}
                    record_getxapi_spend(
                        db              = spend_db,
                        calls_made      = gx_stats.get("calls_made", 0),
                        tweets_collected= gx_stats.get("tweets_collected", result.get("total_tweets", 0)),
                        keyword         = keyword,
                        task_id         = task_id,
                    )
                elif scraper == "instagram":
                    apify_stats = result.get("_apify_run_stats") or {}
                    run_meta = {
                        "usageTotalUsd": apify_stats.get("usageTotalUsd", 0),
                        "stats": {"computeUnits": apify_stats.get("computeUnits", 0)},
                    }
                    record_apify_spend(
                        db            = spend_db,
                        scraper       = "instagram",
                        service_label = "Apify (Instagram)",
                        operation     = "instagram_scrape",
                        run_result    = run_meta,
                        items_count   = apify_stats.get("items_count", result.get("total_posts", 0)),
                        keyword       = keyword,
                        task_id       = task_id,
                    )
                elif scraper == "reddit":
                    record_reddit_spend(spend_db,
                                        items_count=_scraped_counts.get("total_posts", 0),
                                        keyword=keyword, task_id=task_id)
                elif scraper == "autodesk":
                    record_autodesk_spend(spend_db, items_count=_scraped_counts.get("total_posts", 0),
                                          keyword=keyword, task_id=task_id)
                elif scraper == "tiktok":
                    record_scrapecreators_spend(spend_db, items_count=_scraped_counts.get("total_posts", 0),
                                               keyword=keyword, task_id=task_id)
                elif scraper == "edugeek":
                    items = _scraped_counts.get("total_items", 0)
                    record_scrapingbee_spend(spend_db, pages_fetched=max(items * 2, 1),
                                            keyword=keyword, task_id=task_id)
                elif scraper == "spiceworks":
                    from services.spending_service import record_spiceworks_spend
                    record_spiceworks_spend(spend_db, items_count=_scraped_counts.get("total_items", 0),
                                            keyword=keyword, task_id=task_id)
                elif scraper == "quora":
                    from services.spending_service import record_quora_spend
                    record_quora_spend(spend_db, items_count=_scraped_counts.get("total_items", 0),
                                       keyword=keyword, task_id=task_id)
                elif scraper == "facebook":
                    from services.spending_service import record_facebook_spend
                    record_facebook_spend(spend_db,
                                          items_count=result.get("_api_fetched_count")
                                                      or _scraped_counts.get("total_posts", 0),
                                          keyword=keyword, task_id=task_id)
            except Exception as exc:
                logger.warning("Spend recording failed for %s: %s", scraper, exc)

        finished_at = datetime.now(tz=timezone.utc)
        keyword = getattr(cfg, "keyword", "") or ""
        if not keyword and hasattr(cfg, "keywords"):
            keyword = (getattr(cfg, "keywords", None) or [""])[0]
        items = (
            result.get("total_posts") or result.get("total_tweets") or
            result.get("total_articles") or result.get("total_questions") or
            result.get("total_items") or 0
        )
        task_registry[task_id].update({
            "status":      "completed",
            "finished_at": finished_at.isoformat(),
            "result":      result,
        })
        # Update DB
        try:
            if spend_db:
                from db_models import TaskHistory
                row = spend_db.query(TaskHistory).filter_by(task_id=task_id).first()
                if row:
                    row.status      = "completed"
                    row.finished_at = finished_at
                    row.keyword     = keyword[:255] if keyword else None
                    row.items_count = items
                    spend_db.commit()
        except Exception as exc:
            logger.warning("Could not update task in DB: %s", exc)

        scraper_status[scraper].update({
            "last_run":         datetime.now(tz=timezone.utc).isoformat(),
            "last_file":        result.get("file"),
            "total_runs":       scraper_status[scraper]["total_runs"] + 1,
            "last_total_items": items,
        })
        if scraper == "google_news":
            scraper_status[scraper]["last_newsletters_created"] = result.get("newsletters_created", 0)
        logger.info("Task %s (%s) completed", task_id[:8], scraper)

    except Exception as exc:
        logger.exception("Task %s (%s) FAILED: %s", task_id[:8], scraper, exc)
        failed_at = datetime.now(tz=timezone.utc)
        task_registry[task_id].update({
            "status":      "failed",
            "finished_at": failed_at.isoformat(),
            "error":       str(exc),
        })
        # Update DB
        try:
            if spend_db:
                from db_models import TaskHistory
                row = spend_db.query(TaskHistory).filter_by(task_id=task_id).first()
                if row:
                    row.status      = "failed"
                    row.finished_at = failed_at
                    row.error       = str(exc)[:500]
                    spend_db.commit()
        except Exception as exc2:
            logger.warning("Could not update failed task in DB: %s", exc2)
    finally:
        scraper_status[scraper]["running"] = False
        if spend_db is not None:
            try:
                spend_db.close()
            except Exception:
                pass

def _make_task(scraper: str) -> str:
    tid = uuid.uuid4().hex
    now = datetime.now(tz=timezone.utc)
    task_registry[tid] = {
        "task_id":     tid,
        "scraper":     scraper,
        "status":      "queued",
        "started_at":  now.isoformat(),
        "finished_at": None,
        "result":      None,
        "error":       None,
    }
    # Persist to DB
    try:
        db = database.SessionLocal() if database.SessionLocal else None
        if db:
            from db_models import TaskHistory
            db.add(TaskHistory(
                task_id     = tid,
                scraper     = scraper,
                status      = "queued",
                started_at  = now,
                finished_at = None,
                keyword     = None,
                items_count = 0,
                error       = None,
            ))
            db.commit()
            db.close()
    except Exception as exc:
        logger.warning("Could not save task to DB: %s", exc)
    return tid


# ══════════════════════════════════════════════════════════════════════════════
#  Health
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/", tags=["Health"])
def root():
    return {
        "status": "ok", "version": "4.2.0",
        "scrapers": sorted(VALID_SCRAPERS),
        "db_enabled": database.SessionLocal is not None,
        "docs": "/docs",
    }

@app.get("/api/health", tags=["Health"])
def health():
    return {
        "status": "ok",
        "time": datetime.now(tz=timezone.utc).isoformat(),
        "db_enabled": database.SessionLocal is not None,
    }


# ══════════════════════════════════════════════════════════════════════════════
#  Run — individual shortcuts
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/run", response_model=RunResponse, tags=["Run"],
          summary="Run any combination of scrapers in one request")
def run_all(body: RunRequest, background_tasks: BackgroundTasks):
    task_ids: List[str] = []
    for name, cfg in [
        ("reddit", body.reddit),
        ("edugeek", body.edugeek), ("stackexchange", body.stackexchange),
        ("autodesk", body.autodesk), ("twitter", body.twitter),
        ("google_news", body.google_news),
        ("spiceworks", body.spiceworks), ("quora", body.quora),
    ]:
        if cfg is not None:
            tid = _make_task(name)
            background_tasks.add_task(_run_scraper, tid, name, cfg)
            task_ids.append(tid)
    if not task_ids:
        raise HTTPException(400, "No scraper config provided.")
    return RunResponse(message=f"{len(task_ids)} scraper(s) queued.", task_ids=task_ids)


@app.post("/api/run/reddit",        tags=["Run"])
def run_reddit(cfg: RedditConfig, background_tasks: BackgroundTasks):
    tid = _make_task("reddit")
    background_tasks.add_task(_run_scraper, tid, "reddit", cfg)
    return {"message": "Reddit scraper queued.", "task_id": tid}

@app.post("/api/run/tiktok",        tags=["Run"])
def run_tiktok(cfg: TikTokConfig, background_tasks: BackgroundTasks):
    tid = _make_task("tiktok")
    background_tasks.add_task(_run_scraper, tid, "tiktok", cfg)
    return {"message": "TikTok scraper queued.", "task_id": tid}

@app.post("/api/run/edugeek",       tags=["Run"])
def run_edugeek(cfg: EduGeekConfig, background_tasks: BackgroundTasks):
    tid = _make_task("edugeek")
    background_tasks.add_task(_run_scraper, tid, "edugeek", cfg)
    return {"message": "EduGeek scraper queued.", "task_id": tid}

@app.post("/api/run/stackexchange", tags=["Run"])
def run_stackexchange(cfg: StackExchangeConfig, background_tasks: BackgroundTasks):
    tid = _make_task("stackexchange")
    background_tasks.add_task(_run_scraper, tid, "stackexchange", cfg)
    return {"message": "StackExchange scraper queued.", "task_id": tid}

@app.post("/api/run/autodesk",      tags=["Run"])
def run_autodesk(cfg: AutodeskConfig, background_tasks: BackgroundTasks):
    tid = _make_task("autodesk")
    background_tasks.add_task(_run_scraper, tid, "autodesk", cfg)
    return {"message": "Autodesk Community scraper queued.", "task_id": tid}

@app.post("/api/run/twitter",       tags=["Run"])
def run_twitter(cfg: TwitterConfig, background_tasks: BackgroundTasks):
    tid = _make_task("twitter")
    background_tasks.add_task(_run_scraper, tid, "twitter", cfg)
    return {"message": "Twitter scraper queued.", "task_id": tid}

@app.post("/api/run/instagram",     tags=["Run"])
def run_instagram(cfg: InstagramConfig, background_tasks: BackgroundTasks):
    tid = _make_task("instagram")
    background_tasks.add_task(_run_scraper, tid, "instagram", cfg)
    return {"message": "Instagram scraper queued.", "task_id": tid}

@app.post("/api/run/google-news",   tags=["Run"])
def run_google_news(cfg: GoogleNewsConfig, background_tasks: BackgroundTasks):
    tid = _make_task("google_news")
    background_tasks.add_task(_run_scraper, tid, "google_news", cfg)
    return {"message": "Google News scraper queued.", "task_id": tid}

@app.post("/api/run/spiceworks",    tags=["Run"])
def run_spiceworks(cfg: SpiceworksConfig, background_tasks: BackgroundTasks):
    tid = _make_task("spiceworks")
    background_tasks.add_task(_run_scraper, tid, "spiceworks", cfg)
    return {"message": "Spiceworks scraper queued.", "task_id": tid}

@app.post("/api/run/quora",         tags=["Run"])
def run_quora(cfg: QuoraConfig, background_tasks: BackgroundTasks):
    tid = _make_task("quora")
    background_tasks.add_task(_run_scraper, tid, "quora", cfg)
    return {"message": "Quora scraper queued.", "task_id": tid}

@app.post("/api/run/facebook",      tags=["Run"])
def run_facebook(cfg: FacebookConfig, background_tasks: BackgroundTasks):
    tid = _make_task("facebook")
    background_tasks.add_task(_run_scraper, tid, "facebook", cfg)
    return {"message": "Facebook Groups scraper queued.", "task_id": tid}


# ══════════════════════════════════════════════════════════════════════════════
#  Tasks
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/tasks", tags=["Tasks"])
def list_tasks():
    # Merge DB history (last 50) with live in-memory tasks
    tasks = list(task_registry.values())
    try:
        db = database.SessionLocal() if database.SessionLocal else None
        if db:
            from db_models import TaskHistory
            rows = db.query(TaskHistory).order_by(TaskHistory.started_at.desc()).limit(50).all()
            db.close()
            live_ids = {t["task_id"] for t in tasks}
            for row in rows:
                if row.task_id not in live_ids:
                    tasks.append({
                        "task_id":     row.task_id,
                        "scraper":     row.scraper,
                        "status":      row.status,
                        "started_at":  row.started_at.isoformat() if row.started_at else None,
                        "finished_at": row.finished_at.isoformat() if row.finished_at else None,
                        "result":      {
                            "keywords":   [row.keyword] if row.keyword else [],
                            "total_items": row.items_count or 0,
                        },
                        "error":       row.error,
                    })
    except Exception as exc:
        logger.warning("Could not load task history from DB: %s", exc)
    return {
        "total":     len(tasks),
        "queued":    sum(1 for t in tasks if t["status"] == "queued"),
        "running":   sum(1 for t in tasks if t["status"] == "running"),
        "completed": sum(1 for t in tasks if t["status"] == "completed"),
        "failed":    sum(1 for t in tasks if t["status"] == "failed"),
        "tasks":     tasks,
    }

@app.get("/api/tasks/{task_id}", tags=["Tasks"])
def get_task(task_id: str = FPath(...)):
    task = task_registry.get(task_id)
    if not task:
        raise HTTPException(404, f"Task '{task_id}' not found.")
    return task

@app.delete("/api/tasks/{task_id}", tags=["Tasks"])
def delete_task(task_id: str = FPath(...)):
    if task_id not in task_registry:
        raise HTTPException(404, f"Task '{task_id}' not found.")
    if task_registry[task_id]["status"] == "running":
        raise HTTPException(409, "Cannot delete a running task.")
    del task_registry[task_id]
    return {"message": f"Task {task_id} deleted."}


# ══════════════════════════════════════════════════════════════════════════════
#  Scheduling
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/schedule", tags=["Schedule"])
def create_schedule(body: ScheduleRequest):
    created: List[Dict[str, Any]] = []
    def _register(name, cfg):
        if cfg is None:
            return
        if cfg.schedule == "manual":
            raise HTTPException(400, f"Use /api/run for manual runs.")
        mod    = _get_scraper(name)
        job_id = sched.add_job(
            scraper_name=name, interval_name=cfg.schedule,
            fn=mod.run, cfg_dict=cfg.model_dump(), fn_kwargs={"cfg": cfg},
        )
        created.append({"scraper": name, "job_id": job_id, "schedule": cfg.schedule})
    for name, cfg in [
        ("reddit", body.reddit),
        ("edugeek", body.edugeek), ("stackexchange", body.stackexchange),
        ("autodesk", body.autodesk), ("twitter", body.twitter),
        ("google_news", body.google_news),
        ("spiceworks", body.spiceworks), ("quora", body.quora),
    ]:
        _register(name, cfg)
    if not created:
        raise HTTPException(400, "No scraper config provided.")
    return {"message": f"{len(created)} schedule(s) registered.", "jobs": created}

@app.get("/api/schedule",              tags=["Schedule"])
def list_schedules(): return {"jobs": sched.list_jobs()}

@app.delete("/api/schedule/{job_id}", tags=["Schedule"])
def delete_schedule(job_id: str = FPath(...)):
    if not sched.remove_job(job_id):
        raise HTTPException(404, f"Job '{job_id}' not found.")
    return {"message": f"Job {job_id} removed."}

@app.patch("/api/schedule/{job_id}/pause",  tags=["Schedule"])
def pause_schedule(job_id: str = FPath(...)):
    if not sched.pause_job(job_id):
        raise HTTPException(404, f"Job '{job_id}' not found.")
    return {"message": f"Job {job_id} paused."}

@app.patch("/api/schedule/{job_id}/resume", tags=["Schedule"])
def resume_schedule(job_id: str = FPath(...)):
    if not sched.resume_job(job_id):
        raise HTTPException(404, f"Job '{job_id}' not found.")
    return {"message": f"Job {job_id} resumed."}


# ══════════════════════════════════════════════════════════════════════════════
#  Status
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/status", tags=["Status"])
def get_status():
    return scraper_status


# ══════════════════════════════════════════════════════════════════════════════
#  Stats
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/stats/24h", tags=["Stats"],
         summary="Items scraped in last 24h with 7-day comparison")
def stats_24h(db: Session = Depends(get_db)):
    from db_models import ScrapeRun
    from sqlalchemy import func
    from datetime import timedelta

    now      = datetime.now(tz=timezone.utc)
    h24_ago  = now - timedelta(hours=24)
    week_ago = now - timedelta(days=8)

    today_items = db.query(func.coalesce(func.sum(ScrapeRun.total_items), 0)) \
        .filter(ScrapeRun.scraped_at >= h24_ago).scalar() or 0

    prev_items = db.query(func.coalesce(func.sum(ScrapeRun.total_items), 0)) \
        .filter(ScrapeRun.scraped_at >= week_ago,
                ScrapeRun.scraped_at < h24_ago).scalar() or 0

    daily_avg  = prev_items / 7.0 if prev_items > 0 else 0
    change_pct = None
    if daily_avg > 0:
        change_pct = round(((today_items - daily_avg) / daily_avg) * 100, 1)

    return {
        "total_items":   int(today_items),
        "daily_avg_7d":  round(daily_avg, 1),
        "change_7d_pct": change_pct,
    }


@app.get("/api/stats/monthly", tags=["Stats"],
         summary="Items scraped per month for the last 12 months, broken down by scraper")
def stats_monthly(db: Session = Depends(get_db)):
    from db_models import ScrapeRun
    from sqlalchemy import func
    from datetime import timedelta

    cutoff = datetime.now(tz=timezone.utc) - timedelta(days=365)

    rows = (
        db.query(
            func.to_char(ScrapeRun.scraped_at, "YYYY-MM").label("month"),
            ScrapeRun.scraper,
            func.coalesce(func.sum(ScrapeRun.total_items), 0).label("total"),
        )
        .filter(ScrapeRun.scraped_at >= cutoff)
        .group_by("month", ScrapeRun.scraper)
        .order_by("month")
        .all()
    )

    # Aggregate into {month -> {scraper -> count, _total -> count}}
    month_map: dict = {}
    for row in rows:
        m = row.month
        if m not in month_map:
            month_map[m] = {"month": m, "total": 0}
        month_map[m][row.scraper] = int(row.total)
        month_map[m]["total"]    += int(row.total)

    # Build sorted list with human-readable label
    result = []
    for m, data in sorted(month_map.items()):
        try:
            dt    = datetime.strptime(m, "%Y-%m")
            label = dt.strftime("%b %Y")   # "Jan 2025"
        except Exception:
            label = m
        result.append({**data, "label": label})

    return {"months": result}


# ══════════════════════════════════════════════════════════════════════════════
#  Cost Governance — overall budget
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/spending/summary", tags=["Cost"])
def spending_summary(db: Session = Depends(get_db)):
    from services.spending_service import get_spending_summary
    try:
        return get_spending_summary(db)
    except Exception as exc:
        import traceback
        logger.error("spending_summary error: %s\n%s", exc, traceback.format_exc())
        raise HTTPException(500, f"spending_summary error: {exc}")

@app.get("/api/spending/budget", tags=["Cost"])
def get_budget(db: Session = Depends(get_db)):
    from db_models import UserBudget
    row = db.query(UserBudget).filter_by(id=1).first()
    if not row:
        return {
            "monthly_limit_usd":   float(os.environ.get("DEFAULT_MONTHLY_BUDGET_USD", "1000")),
            "alert_threshold_pct": 80,
        }
    return {
        "monthly_limit_usd":   row.monthly_limit_usd,
        "alert_threshold_pct": row.alert_threshold_pct,
    }

@app.post("/api/spending/budget", tags=["Cost"])
def set_budget(body: Dict[str, Any], db: Session = Depends(get_db)):
    from db_models import UserBudget
    limit     = float(body.get("monthly_limit_usd", 1000.0))
    threshold = int(body.get("alert_threshold_pct", 80))
    if limit <= 0:
        raise HTTPException(400, "monthly_limit_usd must be positive")
    if not (1 <= threshold <= 100):
        raise HTTPException(400, "alert_threshold_pct must be 1–100")
    row = db.query(UserBudget).filter_by(id=1).first()
    if row:
        row.monthly_limit_usd   = limit
        row.alert_threshold_pct = threshold
        row.updated_at          = datetime.now(tz=timezone.utc)
    else:
        db.add(UserBudget(
            id=1, monthly_limit_usd=limit,
            alert_threshold_pct=threshold,
            updated_at=datetime.now(tz=timezone.utc),
        ))
    db.commit()
    return {"status": "ok", "monthly_limit_usd": limit, "alert_threshold_pct": threshold}


# ══════════════════════════════════════════════════════════════════════════════
#  Cost Governance — per-scraper budgets
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/spending/scraper-budgets", tags=["Cost"],
         summary="Get all per-scraper budget allocations with spend status")
def get_scraper_budgets(db: Session = Depends(get_db)):
    from services.spending_service import get_scraper_budget_status
    return get_scraper_budget_status(db)

@app.post("/api/spending/scraper-budgets", tags=["Cost"],
          summary="Save per-scraper budget allocations")
def set_scraper_budgets(body: Dict[str, Any], db: Session = Depends(get_db)):
    from db_models import ScraperBudget, UserBudget

    budgets = body.get("budgets", {})
    if not isinstance(budgets, dict):
        raise HTTPException(400, "'budgets' must be an object mapping scraper → amount")

    # Validate sum does not exceed overall budget
    overall_row    = db.query(UserBudget).filter_by(id=1).first()
    default_budget = float(os.environ.get("DEFAULT_MONTHLY_BUDGET_USD", "1000"))
    overall_budget = overall_row.monthly_limit_usd if overall_row else default_budget
    total_alloc    = sum(float(v or 0) for v in budgets.values())

    if total_alloc > overall_budget + 0.01:
        raise HTTPException(
            400,
            f"Total allocation ${total_alloc:.2f} exceeds the overall monthly budget "
            f"${overall_budget:.2f}. Please reduce allocations or increase the overall budget."
        )

    # Validate no budget is set below what's already been spent this month
    from db_models import ApiSpending
    from sqlalchemy import func as _func
    month_start = datetime.now(tz=timezone.utc).replace(
        day=1, hour=0, minute=0, second=0, microsecond=0
    )
    violations = []
    for scraper, amount in budgets.items():
        if scraper not in VALID_SCRAPERS:
            continue
        amount_f = float(amount or 0)
        if amount_f <= 0:
            continue   # setting to 0 (unallocated) is always allowed
        spent = float(
            db.query(_func.coalesce(_func.sum(ApiSpending.cost_usd), 0))
              .filter(ApiSpending.scraper == scraper,
                      ApiSpending.called_at >= month_start)
              .scalar() or 0.0
        )
        if amount_f < spent:
            violations.append(
                f"{scraper}: cannot set ${amount_f:.2f} — already spent ${spent:.4f} this month"
            )
    if violations:
        raise HTTPException(400, "Budget below current spend: " + "; ".join(violations))

    now = datetime.now(tz=timezone.utc)
    for scraper, amount in budgets.items():
        if scraper not in VALID_SCRAPERS:
            continue
        amount = float(amount or 0)
        row = db.query(ScraperBudget).filter_by(scraper=scraper).first()
        if row:
            row.budget_usd = amount
            row.updated_at = now
        else:
            db.add(ScraperBudget(scraper=scraper, budget_usd=amount, updated_at=now))

    db.commit()
    return {"status": "ok", "budgets": budgets, "total_allocated": round(total_alloc, 2)}

@app.get("/api/spending/scraper-status", tags=["Cost"],
         summary="Per-scraper budget usage and block status (for Scraping page)")
def scraper_budget_status(db: Session = Depends(get_db)):
    from services.spending_service import get_scraper_budget_status
    return {"scrapers": get_scraper_budget_status(db)}

@app.get("/api/spending/cost-config", tags=["Cost"],
         summary="Get per-scraper cost rate configuration")
def get_cost_config(db: Session = Depends(get_db)):
    from db_models import ScraperCostConfig
    rows = db.query(ScraperCostConfig).all()
    config = {r.scraper: {
        "cost_mode":  r.cost_mode,
        "cost_value": r.cost_value,
        "cost_per":   r.cost_per,
    } for r in rows}
    # Fill in defaults for scrapers with no DB row
    APIFY_SCRAPERS = {"instagram"}
    for scraper in VALID_SCRAPERS:
        if scraper not in config:
            config[scraper] = {
                "cost_mode":  "apify_real" if scraper in APIFY_SCRAPERS else "free",
                "cost_value": None,
                "cost_per":   None,
            }
    return config


@app.post("/api/spending/cost-config", tags=["Cost"],
          summary="Save per-scraper cost rate configuration")
def set_cost_config(body: Dict[str, Any], db: Session = Depends(get_db)):
    from db_models import ScraperCostConfig
    now = datetime.now(tz=timezone.utc)
    configs = body.get("configs", {})
    if not isinstance(configs, dict):
        raise HTTPException(400, "'configs' must be an object mapping scraper → config")

    for scraper, cfg in configs.items():
        if scraper not in VALID_SCRAPERS:
            continue
        cost_mode  = cfg.get("cost_mode", "free")
        cost_value = cfg.get("cost_value")
        cost_per   = cfg.get("cost_per")

        if cost_mode not in ("apify_real", "per_item", "per_run", "free"):
            raise HTTPException(400, f"Invalid cost_mode '{cost_mode}' for {scraper}")

        cost_value = float(cost_value) if cost_value not in (None, "") else None
        cost_per   = int(cost_per)     if cost_per   not in (None, "") else None

        row = db.query(ScraperCostConfig).filter_by(scraper=scraper).first()
        if row:
            row.cost_mode  = cost_mode
            row.cost_value = cost_value
            row.cost_per   = cost_per
            row.updated_at = now
        else:
            db.add(ScraperCostConfig(
                scraper    = scraper,
                cost_mode  = cost_mode,
                cost_value = cost_value,
                cost_per   = cost_per,
                updated_at = now,
            ))
    db.commit()
    return {"status": "ok", "saved": list(configs.keys())}

@app.get("/api/spending/cost-config", tags=["Cost"],
         summary="Get per-scraper cost rate configuration")
def get_cost_config(db: Session = Depends(get_db)):
    from db_models import ScraperCostConfig
    rows = db.query(ScraperCostConfig).all()
    config = {r.scraper: {
        "cost_mode":  r.cost_mode,
        "cost_value": r.cost_value,
        "cost_per":   r.cost_per,
    } for r in rows}
    # Fill in defaults for scrapers with no DB row
    APIFY_SCRAPERS = {"instagram"}
    for scraper in VALID_SCRAPERS:
        if scraper not in config:
            config[scraper] = {
                "cost_mode":  "apify_real" if scraper in APIFY_SCRAPERS else "free",
                "cost_value": None,
                "cost_per":   None,
            }
    return config


@app.post("/api/spending/cost-config", tags=["Cost"],
          summary="Save per-scraper cost rate configuration")
def set_cost_config(body: Dict[str, Any], db: Session = Depends(get_db)):
    from db_models import ScraperCostConfig
    now = datetime.now(tz=timezone.utc)
    configs = body.get("configs", {})
    if not isinstance(configs, dict):
        raise HTTPException(400, "'configs' must be an object mapping scraper → config")

    for scraper, cfg in configs.items():
        if scraper not in VALID_SCRAPERS:
            continue
        cost_mode  = cfg.get("cost_mode", "free")
        cost_value = cfg.get("cost_value")
        cost_per   = cfg.get("cost_per")

        if cost_mode not in ("apify_real", "per_item", "per_run", "free"):
            raise HTTPException(400, f"Invalid cost_mode '{cost_mode}' for {scraper}")

        cost_value = float(cost_value) if cost_value not in (None, "") else None
        cost_per   = int(cost_per)     if cost_per   not in (None, "") else None

        row = db.query(ScraperCostConfig).filter_by(scraper=scraper).first()
        if row:
            row.cost_mode  = cost_mode
            row.cost_value = cost_value
            row.cost_per   = cost_per
            row.updated_at = now
        else:
            db.add(ScraperCostConfig(
                scraper    = scraper,
                cost_mode  = cost_mode,
                cost_value = cost_value,
                cost_per   = cost_per,
                updated_at = now,
            ))
    db.commit()
    return {"status": "ok", "saved": list(configs.keys())}


# ══════════════════════════════════════════════════════════════════════════════
#  Cost Governance — history, emails, alerts
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/spending/history", tags=["Cost"])
def spending_history(
    limit:  int = Query(50, le=200),
    offset: int = Query(0),
    db:     Session = Depends(get_db),
):
    from db_models import ApiSpending
    total = db.query(ApiSpending).count()
    rows  = (
        db.query(ApiSpending)
          .order_by(ApiSpending.called_at.desc())
          .offset(offset).limit(limit).all()
    )
    return {
        "total": total,
        "rows": [
            {
                "id": r.id, "provider": r.provider, "service": r.service,
                "operation": r.operation, "scraper": r.scraper, "task_id": r.task_id,
                "cost_usd": r.cost_usd, "cost_units": r.cost_units,
                "is_estimated": r.is_estimated, "items_count": r.items_count,
                "keyword": r.keyword,
                "called_at": r.called_at.isoformat() if r.called_at else None,
            }
            for r in rows
        ],
    }

@app.get("/api/spending/alert-emails", tags=["Cost"])
def get_alert_emails(db: Session = Depends(get_db)):
    from db_models import BudgetAlertEmail
    rows = db.query(BudgetAlertEmail).order_by(BudgetAlertEmail.added_at).all()
    return {"emails": [r.email for r in rows]}

@app.post("/api/spending/alert-emails", tags=["Cost"])
def save_alert_emails(body: Dict[str, Any], db: Session = Depends(get_db)):
    from db_models import BudgetAlertEmail
    emails = body.get("emails", [])
    if not isinstance(emails, list):
        raise HTTPException(400, "emails must be a list")
    db.query(BudgetAlertEmail).delete()
    for email in emails:
        email = email.strip().lower()
        if email:
            db.add(BudgetAlertEmail(email=email, added_at=datetime.now(tz=timezone.utc)))
    db.commit()
    return {"status": "ok", "emails": emails}

@app.post("/api/spending/trigger-alert", tags=["Cost"])
def trigger_budget_alert(body: Dict[str, Any], db: Session = Depends(get_db)):
    from db_models import UserBudget, BudgetAlertEmail
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    budget_pct = float(body.get("budget_pct", 0))
    alert_type = body.get("alert_type", "warning")
    is_blocked = alert_type == "blocked"

    row        = db.query(UserBudget).filter_by(id=1).first()
    budget_usd = row.monthly_limit_usd if row else float(
        os.environ.get("DEFAULT_MONTHLY_BUDGET_USD", "1000")
    )
    emails = [r.email for r in db.query(BudgetAlertEmail).all()]

    if not emails:
        return {"status": "no_recipients"}

    smtp_host = os.environ.get("ALERT_SMTP_HOST", "smtp.gmail.com")
    smtp_port = int(os.environ.get("ALERT_SMTP_PORT", "587"))
    smtp_user = os.environ.get("ALERT_SMTP_USER", "")
    smtp_pass = os.environ.get("ALERT_SMTP_PASS", "")

    if not smtp_user or not smtp_pass:
        return {"status": "smtp_not_configured", "would_have_sent_to": emails}

    color   = "#dc2626" if is_blocked else "#d97706"
    icon    = "🚫" if is_blocked else "⚠️"
    subject = f"{icon} Overall Budget {'BLOCKED' if is_blocked else 'Warning'} — {budget_pct:.1f}% used"
    action  = (
        "All scrapers have been <strong>permanently blocked</strong>. Increase your budget to resume."
        if is_blocked else
        "Scrapers are still running. They will be blocked at 100%."
    )

    html = f"""
    <html><body style="font-family:sans-serif;background:#0a0e17;color:#e2e8f0;padding:24px">
      <div style="max-width:540px;margin:0 auto;background:#111827;border-radius:12px;
                  border:2px solid {color};padding:32px">
        <h2 style="color:{color};margin-top:0">{icon} Budget Alert</h2>
        <p>Overall budget is at <strong style="color:{color}">{budget_pct:.1f}%</strong>
           of <strong>${budget_usd:,.2f}</strong> monthly limit.</p>
        <p>{action}</p>
        <hr style="border-color:#1e293b;margin:24px 0"/>
        <p style="color:#64748b;font-size:0.8rem">
          Manage at <strong>Cost Governance → Modifications</strong>
        </p>
      </div>
    </body></html>
    """

    sent, failed = [], []
    try:
        with smtplib.SMTP(smtp_host, smtp_port, timeout=15) as srv:
            srv.ehlo(); srv.starttls(); srv.ehlo()
            srv.login(smtp_user, smtp_pass)
            for recipient in emails:
                try:
                    msg = MIMEMultipart("alternative")
                    msg["Subject"] = subject
                    msg["From"]    = f"TrendSense Alerts <{smtp_user}>"
                    msg["To"]      = recipient
                    msg.attach(MIMEText(html, "html"))
                    srv.sendmail(smtp_user, recipient, msg.as_string())
                    sent.append(recipient)
                except Exception as exc:
                    failed.append({"email": recipient, "error": str(exc)})
    except smtplib.SMTPAuthenticationError:
        return {"status": "smtp_auth_error",
                "error": "Use a Google App Password, not your Gmail password."}
    except Exception as exc:
        return {"status": "smtp_error", "error": str(exc)}

    return {"status": "ok", "alert_type": alert_type, "sent": sent, "failed": failed}

@app.get("/api/spending/debug", tags=["Cost"])
def spending_debug(db: Session = Depends(get_db)):
    try:
        from db_models import ApiSpending, UserBudget
        rows   = db.query(ApiSpending).order_by(ApiSpending.called_at.desc()).limit(10).all()
        budget = db.query(UserBudget).filter_by(id=1).first()
        return {
            "api_spending_count": db.query(ApiSpending).count(),
            "user_budget": {
                "monthly_limit_usd":   budget.monthly_limit_usd   if budget else None,
                "alert_threshold_pct": budget.alert_threshold_pct if budget else None,
            },
            "last_10_rows": [
                {
                    "id": r.id, "provider": r.provider, "scraper": r.scraper,
                    "cost_usd": r.cost_usd, "keyword": r.keyword,
                    "called_at": r.called_at.isoformat() if r.called_at else None,
                }
                for r in rows
            ],
        }
    except Exception as exc:
        import traceback
        return {"error": str(exc), "traceback": traceback.format_exc()}


# ══════════════════════════════════════════════════════════════════════════════
#  LLM Configuration & Feed
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/llm/configs", tags=["LLM"])
def get_llm_configs(db: Session = Depends(get_db)):
    from llm_service import get_all_configs
    return get_all_configs(db)

@app.get("/api/llm/active-config", tags=["LLM"])
def get_active_llm_config(db: Session = Depends(get_db)):
    from llm_service import get_active_config
    cfg = get_active_config(db)
    if not cfg:
        return {"configured": False}
    return {"configured": True, "provider": cfg["provider"], "model": cfg["model"]}

@app.post("/api/llm/config", tags=["LLM"])
def save_llm_config(body: Dict[str, Any], db: Session = Depends(get_db)):
    from llm_service import save_provider_config
    provider   = body.get("provider", "")
    api_key    = body.get("api_key")
    model      = body.get("model", "")
    set_active = bool(body.get("set_active", False))
    if not provider or not model:
        raise HTTPException(400, "provider and model are required")
    try:
        return save_provider_config(db, provider, api_key, model, set_active)
    except ValueError as e:
        raise HTTPException(400, str(e))

@app.post("/api/llm/enhance-prompt", tags=["LLM"])
def enhance_prompt_endpoint(body: Dict[str, Any], db: Session = Depends(get_db)):
    from llm_service import enhance_prompt
    raw_prompt   = body.get("prompt", "").strip()
    data_sources = body.get("data_sources", [])
    sample_rows  = body.get("sample_rows", [])
    if not raw_prompt:
        raise HTTPException(400, "prompt is required")
    try:
        return enhance_prompt(db, raw_prompt, data_sources, sample_rows)
    except RuntimeError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"Enhancement failed: {e}")

@app.post("/api/llm/feed", tags=["LLM"])
def feed_to_llm_endpoint(body: Dict[str, Any], db: Session = Depends(get_db)):
    from llm_service import feed_to_llm
    prompt  = body.get("prompt", "").strip()
    rows    = body.get("rows", [])
    keyword = body.get("keyword", "")
    if not prompt:
        raise HTTPException(400, "prompt is required")
    if not rows:
        raise HTTPException(400, "rows (data) is required")
    if len(rows) > 15:
        raise HTTPException(400, f"Maximum 15 records allowed per LLM request (got {len(rows)})")
    try:
        return feed_to_llm(db, prompt, rows, keyword)
    except RuntimeError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.exception("feed_to_llm failed: %s", e)
        raise HTTPException(500, f"LLM call failed: {e}")

@app.get("/api/llm/spending", tags=["LLM"])
def llm_spending_summary(db: Session = Depends(get_db)):
    from llm_service import get_llm_spending_summary
    return get_llm_spending_summary(db)

@app.post("/api/llm/analyses", tags=["LLM"])
def save_llm_analysis(body: Dict[str, Any], db: Session = Depends(get_db)):
    """Save a completed LLM analysis to the DB."""
    import json as _json
    from db_models import LLMAnalysis
    generated_at_raw = body.get("generatedAt") or body.get("generated_at")
    if not generated_at_raw:
        raise HTTPException(400, "generatedAt is required")
    try:
        generated_at = datetime.fromisoformat(
            str(generated_at_raw).replace("Z", "+00:00")
        )
    except ValueError:
        raise HTTPException(400, f"Invalid generatedAt: {generated_at_raw!r}")

    platforms = body.get("platforms", [])
    row = LLMAnalysis(
        provider        = body.get("provider", ""),
        model           = body.get("model", ""),
        raw_prompt      = body.get("rawPrompt", ""),
        enhanced_prompt = body.get("enhancedPrompt", ""),
        response        = body.get("response", ""),
        record_count    = int(body.get("recordCount", 0) or 0),
        tokens_used     = int(body.get("tokens_used", 0) or 0),
        cost_usd        = float(body.get("cost_usd", 0.0) or 0.0),
        platforms       = _json.dumps(platforms) if isinstance(platforms, list) else str(platforms),
        generated_at    = generated_at,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    logger.info("LLM analysis saved: id=%d provider=%s", row.id, row.provider)
    return {"id": row.id}


@app.get("/api/llm/analyses", tags=["LLM"])
def list_llm_analyses(
    date: str = Query(None, description="Filter by date YYYY-MM-DD"),
    limit: int = Query(100, ge=1, le=500),
    db: Session = Depends(get_db),
):
    """Return saved LLM analyses, newest first. Optional ?date=YYYY-MM-DD filter."""
    import json as _json
    from db_models import LLMAnalysis
    q = db.query(LLMAnalysis)
    if date:
        try:
            day_start = datetime.fromisoformat(f"{date}T00:00:00+00:00")
            day_end   = datetime.fromisoformat(f"{date}T23:59:59+00:00")
            q = q.filter(LLMAnalysis.generated_at >= day_start,
                         LLMAnalysis.generated_at <= day_end)
        except ValueError:
            raise HTTPException(400, f"Invalid date: {date!r}")
    rows = q.order_by(LLMAnalysis.generated_at.desc()).limit(limit).all()

    def _row(r):
        try:
            platforms = _json.loads(r.platforms) if r.platforms else []
        except Exception:
            platforms = []
        return {
            "id":              r.id,
            "provider":        r.provider,
            "model":           r.model,
            "rawPrompt":       r.raw_prompt,
            "enhancedPrompt":  r.enhanced_prompt,
            "response":        r.response,
            "recordCount":     r.record_count,
            "tokens_used":     r.tokens_used,
            "cost_usd":        r.cost_usd,
            "platforms":       platforms,
            "generatedAt":     r.generated_at.isoformat() if r.generated_at else None,
        }

    return {"total": len(rows), "analyses": [_row(r) for r in rows]}


@app.delete("/api/llm/analyses/{analysis_id}", tags=["LLM"])
def delete_llm_analysis(analysis_id: int = FPath(...), db: Session = Depends(get_db)):
    """Delete a saved LLM analysis by id."""
    from db_models import LLMAnalysis
    row = db.query(LLMAnalysis).filter_by(id=analysis_id).first()
    if not row:
        raise HTTPException(404, f"Analysis {analysis_id} not found")
    db.delete(row)
    db.commit()
    return {"status": "ok", "deleted_id": analysis_id}


@app.delete("/api/llm/config/{provider}/key", tags=["LLM"])
def delete_llm_key(provider: str = FPath(...), db: Session = Depends(get_db)):
    from db_models import LLMProviderConfig
    if provider not in ("openai", "anthropic", "gemini"):
        raise HTTPException(400, f"Unknown provider: {provider}")
    row = db.query(LLMProviderConfig).filter_by(provider=provider).first()
    if not row:
        raise HTTPException(404, "Provider not configured")
    row.api_key   = None
    row.is_active = False
    db.commit()
    return {"status": "ok", "provider": provider, "message": "API key removed"}


# ══════════════════════════════════════════════════════════════════════════════
#  Webhooks — Power Automate Adaptive Card response
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/api/webhook/google-news/response", tags=["Webhooks"])
async def webhook_google_news_response(
    request: Request,
    db: Session = Depends(get_db),
):
    """
    Receives the Adaptive Card submission from Power Automate.

    Power Automate HTTP action body (set to JSON):
    {
      "action":     "approve",          ← or "reject"
      "job_id":     "abc123def456...",
      "selected_0": "true",
      "selected_1": "false",
      "selected_2": "true",
      ...
    }
    """
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid JSON body")

    from newsletter_service import handle_teams_submission

    action  = body.get("action")  or request.query_params.get("action", "approve")
    job_id  = body.get("job_id") or request.query_params.get("job_id", "")

    if not job_id:
        raise HTTPException(status_code=400, detail="job_id is required")

    try:
        result = handle_teams_submission(db, {**body, "action": action, "job_id": job_id})
        return {"status": "ok", **result}
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))


# ══════════════════════════════════════════════════════════════════════════════
#  Newsletter & Webhook
# ══════════════════════════════════════════════════════════════════════════════

@app.post("/webhook/google-news/response", tags=["Newsletter"])
def webhook_google_news_response(
    job_id: str = None,
    action: str = None,
    body: Dict[str, Any] = None,
    db: Session = Depends(get_db)
):
    from newsletter_service import process_webhook_response

    # Handle GET request (from Teams Adaptive Card buttons)
    if action:
        if not job_id:
            raise HTTPException(400, "job_id is required")
        approved = action == "approve"
        reason = "Approved via Teams" if approved else "Rejected via Teams"
    else:
        # Handle POST request (existing)
        if body is None:
            raise HTTPException(400, "Request body is required")
        job_id = body.get("job_id", "").strip()
        approved = bool(body.get("approved", False))
        reason = body.get("reason", "")

    if not job_id:
        raise HTTPException(400, "job_id is required")

    try:
        return process_webhook_response(db, job_id, approved, reason)
    except ValueError as e:
        raise HTTPException(404, str(e))
    except RuntimeError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.exception("Webhook response processing failed: %s", e)
        raise HTTPException(500, f"Processing failed: {e}")

@app.get("/api/newsletter/jobs", tags=["Newsletter"])
def get_newsletter_jobs(db: Session = Depends(get_db)):
    from newsletter_service import get_all_jobs
    return {"jobs": get_all_jobs(db)}

@app.get("/api/newsletter/pending", tags=["Newsletter"])
def get_pending_newsletter_jobs(db: Session = Depends(get_db)):
    from newsletter_service import get_pending_jobs
    return {"jobs": get_pending_jobs(db)}

@app.get("/api/newsletters", tags=["Newsletter"])
def get_newsletters(db: Session = Depends(get_db)):
    from newsletter_service import get_all_newsletters
    return {"newsletters": get_all_newsletters(db)}

@app.get("/api/newsletters/{newsletter_id}", tags=["Newsletter"])
def get_newsletter(newsletter_id: int = FPath(...), db: Session = Depends(get_db)):
    from newsletter_service import get_newsletter_by_id
    nl = get_newsletter_by_id(db, newsletter_id)
    if not nl:
        raise HTTPException(404, f"Newsletter {newsletter_id} not found")
    return nl

@app.delete("/api/newsletters/{newsletter_id}", tags=["Newsletter"])
def delete_newsletter(newsletter_id: int = FPath(...), db: Session = Depends(get_db)):
    from db_models import GeneratedNewsletter
    nl = db.query(GeneratedNewsletter).filter(GeneratedNewsletter.id == newsletter_id).first()
    if not nl:
        raise HTTPException(404, f"Newsletter {newsletter_id} not found")
    db.delete(nl)
    db.commit()
    return {"deleted": newsletter_id}


# ══════════════════════════════════════════════════════════════════════════════
#  Keywords
# ══════════════════════════════════════════════════════════════════════════════

class _KeywordAddBody(BaseModel):
    keywords: List[str]
    pool:     str = "shared"  # "shared" | "google_news"

@app.get("/api/keywords", tags=["Keywords"])
def list_keywords(db: Session = Depends(get_db)):
    from db_models import ScraperKeyword
    rows = db.query(ScraperKeyword).order_by(ScraperKeyword.created_at).all()
    return {
        "shared":      [{"id": r.id, "keyword": r.keyword} for r in rows if r.pool == "shared"],
        "google_news": [{"id": r.id, "keyword": r.keyword} for r in rows if r.pool == "google_news"],
    }

@app.post("/api/keywords", tags=["Keywords"])
def add_keywords(body: _KeywordAddBody, db: Session = Depends(get_db)):
    from db_models import ScraperKeyword
    added = []
    for kw in body.keywords:
        kw = kw.strip()
        if not kw:
            continue
        exists = db.query(ScraperKeyword).filter(
            ScraperKeyword.keyword == kw,
            ScraperKeyword.pool    == body.pool,
        ).first()
        if not exists:
            db.add(ScraperKeyword(
                keyword=kw, pool=body.pool,
                created_at=datetime.now(tz=timezone.utc),
            ))
            added.append(kw)
    db.commit()
    return {"added": added}

@app.delete("/api/keywords/{keyword_id}", tags=["Keywords"])
def delete_keyword(keyword_id: int = FPath(...), db: Session = Depends(get_db)):
    from db_models import ScraperKeyword
    kw = db.query(ScraperKeyword).filter(ScraperKeyword.id == keyword_id).first()
    if not kw:
        raise HTTPException(404, f"Keyword {keyword_id} not found")
    db.delete(kw)
    db.commit()
    return {"deleted": keyword_id}


# ══════════════════════════════════════════════════════════════════════════════
#  Facebook Groups
# ══════════════════════════════════════════════════════════════════════════════

class _FacebookGroupBody(BaseModel):
    name: str
    url:  str

@app.get("/api/facebook/groups", tags=["Facebook Groups"])
def list_facebook_groups(db: Session = Depends(get_db)):
    from db_models import FacebookGroup
    rows = db.query(FacebookGroup).order_by(FacebookGroup.created_at).all()
    return {"groups": [{"id": r.id, "name": r.name, "url": r.url} for r in rows]}

@app.post("/api/facebook/groups", tags=["Facebook Groups"])
def add_facebook_group(body: _FacebookGroupBody, db: Session = Depends(get_db)):
    from db_models import FacebookGroup
    name = body.name.strip()
    url  = body.url.strip()
    if not name or not url:
        raise HTTPException(400, "name and url are required")
    existing = db.query(FacebookGroup).filter(FacebookGroup.url == url).first()
    if existing:
        raise HTTPException(409, "A group with this URL already exists")
    grp = FacebookGroup(name=name, url=url, created_at=datetime.now(tz=timezone.utc))
    db.add(grp)
    db.commit()
    db.refresh(grp)
    return {"id": grp.id, "name": grp.name, "url": grp.url}

@app.delete("/api/facebook/groups/{group_id}", tags=["Facebook Groups"])
def delete_facebook_group(group_id: int = FPath(...), db: Session = Depends(get_db)):
    from db_models import FacebookGroup
    grp = db.query(FacebookGroup).filter(FacebookGroup.id == group_id).first()
    if not grp:
        raise HTTPException(404, f"Group {group_id} not found")
    db.delete(grp)
    db.commit()
    return {"deleted": group_id}


# ══════════════════════════════════════════════════════════════════════════════
#  Keyword Selections
# ══════════════════════════════════════════════════════════════════════════════

class _KwSelectionBody(BaseModel):
    scraper:    str
    keyword_id: int

@app.get("/api/keyword-selections", tags=["Keywords"])
def get_keyword_selections(db: Session = Depends(get_db)):
    """Return { scraper: [keyword_id, ...] } for all scrapers."""
    from db_models import ScraperKeywordSelection
    rows = db.query(ScraperKeywordSelection).all()
    result: dict[str, list[int]] = {}
    for r in rows:
        result.setdefault(r.scraper, []).append(r.keyword_id)
    return {"selections": result}

@app.post("/api/keyword-selections", tags=["Keywords"])
def add_keyword_selection(body: _KwSelectionBody, db: Session = Depends(get_db)):
    """Select a keyword for a scraper (idempotent)."""
    from db_models import ScraperKeywordSelection
    existing = db.query(ScraperKeywordSelection).filter_by(
        scraper=body.scraper, keyword_id=body.keyword_id
    ).first()
    if not existing:
        db.add(ScraperKeywordSelection(scraper=body.scraper, keyword_id=body.keyword_id))
        db.commit()
    return {"scraper": body.scraper, "keyword_id": body.keyword_id}

@app.delete("/api/keyword-selections", tags=["Keywords"])
def remove_keyword_selection(body: _KwSelectionBody, db: Session = Depends(get_db)):
    """Deselect a keyword for a scraper."""
    from db_models import ScraperKeywordSelection
    row = db.query(ScraperKeywordSelection).filter_by(
        scraper=body.scraper, keyword_id=body.keyword_id
    ).first()
    if row:
        db.delete(row)
        db.commit()
    return {"scraper": body.scraper, "keyword_id": body.keyword_id}


# ══════════════════════════════════════════════════════════════════════════════
#  Smart Brain
# ══════════════════════════════════════════════════════════════════════════════

_LLM_STRIP_FIELDS: dict[str, set] = {
    "reddit":        {"score"},
    "tiktok":        {"plays", "likes"},
    "autodesk":      {"kudos"},
    "stackexchange": {"score"},
    "instagram":     {"likes", "comments_count"},
    "twitter":       {"retweets", "replies", "avatar"},
    "quora":         {"answer_count", "body"},
}

def _smart_brain_records_for_runs(db, runs, max_per_run: int = 100) -> list[dict]:
    """Fetch preview records for a list of ScrapeRun ORM objects."""
    from services.search_service import (
        _reddit_post_preview, _tiktok_post_preview, _edugeek_post_preview,
        _autodesk_post_preview, _se_question_preview, _gnews_preview,
        _instagram_post_preview, _spiceworks_preview, _twitter_tweet_preview,
        _quora_question_preview,
    )
    from db_models import (
        RedditPost, TikTokPost, EduGeekPost, AutodeskPost, StackExchangeQuestion,
        GoogleNewsArticle, InstagramPost, SpiceworksPost, TwitterTweet, QuoraQuestion,
    )
    SOURCE_MAP = {
        "reddit":        (RedditPost,             _reddit_post_preview),
        "tiktok":        (TikTokPost,             _tiktok_post_preview),
        "edugeek":       (EduGeekPost,            _edugeek_post_preview),
        "autodesk":      (AutodeskPost,           _autodesk_post_preview),
        "stackexchange": (StackExchangeQuestion,  _se_question_preview),
        "instagram":     (InstagramPost,          _instagram_post_preview),
        "spiceworks":    (SpiceworksPost,         _spiceworks_preview),
        "twitter":       (TwitterTweet,           _twitter_tweet_preview),
        "quora":         (QuoraQuestion,          _quora_question_preview),
    }
    rows = []
    for run in runs:
        scraper = (run.scraper or "").lower()
        if scraper == "google_news":
            q = db.query(GoogleNewsArticle)
            if run.keyword:
                q = q.filter(GoogleNewsArticle.search_query == run.keyword)
            elif run.scraped_at:
                window_start = run.scraped_at - timedelta(minutes=10)
                window_end   = run.scraped_at + timedelta(minutes=10)
                q = q.filter(
                    GoogleNewsArticle.scraped_at >= window_start,
                    GoogleNewsArticle.scraped_at <= window_end,
                )
            recs = q.limit(max_per_run).all()
            rows.extend(_gnews_preview(r, db) for r in recs)
            continue
        entry = SOURCE_MAP.get(scraper)
        if not entry:
            continue
        model, preview_fn = entry
        try:
            recs   = db.query(model).filter(model.run_id == run.id).limit(max_per_run).all()
            strip  = _LLM_STRIP_FIELDS.get(scraper, set())
            for r in recs:
                row = preview_fn(r, db)
                for f in strip:
                    row.pop(f, None)
                rows.append(row)
        except Exception as exc:
            logger.warning("smart_brain: failed to fetch %s run %s: %s", scraper, run.id, exc)
    return rows


@app.get("/api/smart-brain/sessions", tags=["SmartBrain"])
def smart_brain_sessions(
    from_date: str | None = Query(None),
    limit:     int        = Query(10, le=500),
    order:     str        = Query("desc"),      # "asc" | "desc"
    db:        Session    = Depends(get_db),
):
    from db_models import ScrapeRun
    # "From Date" mode uses ASC so sessions start at the entered date going forward.
    # "Last Session" mode uses DESC so the newest session is first.
    asc_order = order == "asc"
    q = db.query(ScrapeRun)
    if from_date:
        try:
            dt = datetime.fromisoformat(from_date)
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            q = q.filter(ScrapeRun.scraped_at >= dt)
        except Exception:
            pass
    q = q.order_by(ScrapeRun.scraped_at.asc() if asc_order else ScrapeRun.scraped_at.desc())
    runs = q.all()

    # Group consecutive runs within 10 minutes into the same session.
    # abs() handles both ASC and DESC run ordering.
    sessions: list[list] = []
    current:  list       = []
    prev_dt              = None
    for run in runs:
        run_dt = run.scraped_at
        if run_dt and run_dt.tzinfo is None:
            run_dt = run_dt.replace(tzinfo=timezone.utc)
        if prev_dt is None or abs((prev_dt - (run_dt or prev_dt)).total_seconds()) <= 600:
            current.append(run)
            prev_dt = run_dt or prev_dt
        else:
            sessions.append(current)
            current = [run]
            prev_dt = run_dt
    if current:
        sessions.append(current)

    sessions = sessions[:limit]
    result = []
    for i, grp in enumerate(sessions):
        first    = grp[0]
        first_dt = first.scraped_at
        result.append({
            "session_index": i,
            "scraped_at":    first_dt.isoformat() if first_dt else None,
            "scraper_count": len(grp),
            "total_items":   sum(r.total_items or 0 for r in grp),
            "runs": [
                {
                    "run_id":      r.id,
                    "scraper":     r.scraper,
                    "keyword":     r.keyword or "",
                    "scraped_at":  r.scraped_at.isoformat() if r.scraped_at else None,
                    "total_items": r.total_items or 0,
                }
                for r in grp
            ],
        })
    return {"sessions": result}


@app.post("/api/smart-brain/parse-file", tags=["SmartBrain"])
async def smart_brain_parse_file(file: UploadFile = File(...)):
    content  = await file.read()
    filename = file.filename or ""
    ext      = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    text     = ""
    try:
        if ext == "txt":
            text = content.decode("utf-8", errors="replace")
        elif ext == "pdf":
            import io
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text   = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in ("doc", "docx"):
            import io
            from docx import Document
            doc  = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            raise HTTPException(400, "Unsupported file type. Upload a .txt, .pdf, or .docx file.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, f"Failed to parse file: {exc}")
    return {"text": text.strip(), "filename": filename}


@app.post("/api/smart-brain/run", tags=["SmartBrain"])
def smart_brain_run(
    body: SmartBrainRunRequest,
    db:   Session = Depends(get_db),
):
    if not body.prompt.strip():
        raise HTTPException(400, "Prompt must not be empty.")
    if not body.run_ids:
        raise HTTPException(400, "Select at least one scraper run.")

    from db_models import ScrapeRun
    runs = db.query(ScrapeRun).filter(ScrapeRun.id.in_(body.run_ids)).all()
    if not runs:
        raise HTTPException(404, "None of the selected run IDs were found.")

    data_rows = _smart_brain_records_for_runs(db, runs, max_per_run=max(1, body.max_per_run))
    if not data_rows:
        raise HTTPException(400, "No records found for the selected runs.")

    from llm_service import feed_to_llm, enhance_prompt
    try:
        enhanced = enhance_prompt(db, body.prompt, sample_rows=data_rows[:3])
        ep       = enhanced.get("enhanced_prompt", body.prompt)
    except Exception:
        ep = body.prompt

    try:
        result = feed_to_llm(db, ep, data_rows, keyword=body.keyword or body.prompt[:80])
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))

    return {
        "response":     result["response"],
        "provider":     result["provider"],
        "model":        result["model"],
        "tokens_used":  result.get("tokens_used", 0),
        "cost_usd":     result.get("cost_usd", 0.0),
        "records_used": len(data_rows),
    }


@app.post("/api/smart-brain/export-docx", tags=["SmartBrain"])
def smart_brain_export_docx(body: SmartBrainExportRequest):
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    doc = Document()
    doc.add_heading(body.title, level=0)

    for line in body.text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 40)
        else:
            p = doc.add_paragraph()
            parts = stripped.split("**")
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="smart_brain_analysis.docx"'},
    )


@app.post("/api/smart-brain/enhance-single", tags=["SmartBrain"])
def smart_brain_enhance_single(
    body: SmartBrainEnhanceSingleRequest,
    db:   Session = Depends(get_db),
):
    if not body.prompt.strip():
        raise HTTPException(400, "Prompt must not be empty.")
    from llm_service import enhance_prompt
    try:
        result = enhance_prompt(db, body.prompt, sample_rows=[body.record])
        return {
            "enhanced_prompt":  result.get("enhanced_prompt", body.prompt),
            "summary_for_user": result.get("summary_for_user", ""),
        }
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))


@app.post("/api/smart-brain/run-direct", tags=["SmartBrain"])
def smart_brain_run_direct(
    body: SmartBrainDirectRunRequest,
    db:   Session = Depends(get_db),
):
    if not body.prompt.strip():
        raise HTTPException(400, "Prompt must not be empty.")
    if not body.records:
        raise HTTPException(400, "No records provided.")
    cleaned = []
    for rec in body.records:
        source = rec.get("source", "")
        strip  = _LLM_STRIP_FIELDS.get(source, set())
        cleaned.append({k: v for k, v in rec.items() if k not in strip})
    from llm_service import feed_to_llm
    try:
        result = feed_to_llm(db, body.prompt, cleaned)
    except RuntimeError as exc:
        raise HTTPException(400, str(exc))
    return {
        "response":     result["response"],
        "provider":     result["provider"],
        "model":        result["model"],
        "tokens_used":  result.get("tokens_used", 0),
        "cost_usd":     result.get("cost_usd", 0.0),
        "records_used": len(cleaned),
    }


@app.post("/api/smart-brain/history", tags=["SmartBrain"])
def save_smart_brain_analysis(body: dict = Body(...), db: Session = Depends(get_db)):
    from db_models import SmartBrainAnalysis
    entry = SmartBrainAnalysis(
        result          = body.get("result", ""),
        provider        = body.get("provider", ""),
        model           = body.get("model", ""),
        tokens_used     = body.get("tokens_used", 0),
        cost_usd        = body.get("cost_usd", 0.0),
        enhanced_prompt = body.get("enhanced_prompt", ""),
        prompt_used     = body.get("prompt_used", ""),
        record_count    = body.get("record_count", 0),
    )
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return {"id": entry.id, "created_at": entry.created_at.isoformat()}


@app.get("/api/smart-brain/history", tags=["SmartBrain"])
def get_smart_brain_history(limit: int = Query(100, le=200), db: Session = Depends(get_db)):
    from db_models import SmartBrainAnalysis
    rows = db.query(SmartBrainAnalysis).order_by(SmartBrainAnalysis.created_at.desc()).limit(limit).all()
    return {"analyses": [
        {
            "id":              r.id,
            "result":          r.result,
            "provider":        r.provider,
            "model":           r.model,
            "tokens_used":     r.tokens_used,
            "cost_usd":        r.cost_usd,
            "enhanced_prompt": r.enhanced_prompt,
            "prompt_used":     r.prompt_used,
            "record_count":    r.record_count,
            "timestamp":       r.created_at.isoformat() if r.created_at else "",
        }
        for r in rows
    ]}


@app.delete("/api/smart-brain/history/{entry_id}", tags=["SmartBrain"])
def delete_smart_brain_analysis(entry_id: int, db: Session = Depends(get_db)):
    from db_models import SmartBrainAnalysis
    entry = db.query(SmartBrainAnalysis).filter_by(id=entry_id).first()
    if not entry:
        raise HTTPException(404, "Analysis not found.")
    db.delete(entry)
    db.commit()
    return {"deleted": entry_id}


# ══════════════════════════════════════════════════════════════════════════════
#  Search & Export
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/keywords/used", tags=["Search"])
def get_used_keywords(db: Session = Depends(get_db)):
    from db_models import ScrapeRun, GoogleNewsArticle
    run_kws   = [r[0] for r in db.query(ScrapeRun.keyword).filter(
        ScrapeRun.keyword.isnot(None), ScrapeRun.keyword != ""
    ).all()]
    gnews_kws = [r[0] for r in db.query(GoogleNewsArticle.search_query).filter(
        GoogleNewsArticle.search_query.isnot(None), GoogleNewsArticle.search_query != ""
    ).all()]
    # Deduplicate case-insensitively — keep lowercase version
    seen: dict = {}
    for kw in run_kws + gnews_kws:
        seen[kw.lower().strip()] = kw.lower().strip()
    all_kws = sorted(seen.values(), key=str.lower)
    return {"keywords": all_kws}


@app.get("/search", tags=["Search"])
def search(
    keyword:         str | None = Query(None),
    source:          str | None = Query(None),
    limit:           int        = Query(50, le=100),
    offset:          int        = Query(0),
    date_range:      str | None = Query(None),
    scrape_keyword:  str | None = Query(None),
    group_url:       str | None = Query(None),
    db:              Session    = Depends(get_db),
):
    from services.search_service import SearchService, ALL_SOURCES
    if source and source not in ALL_SOURCES:
        raise HTTPException(400, f"Unknown source '{source}'. Valid: {ALL_SOURCES}")
    sk  = scrape_keyword.strip() if scrape_keyword and scrape_keyword.strip() else None
    gu  = group_url.strip()      if group_url      and group_url.strip()      else None
    if keyword and len(keyword.strip()) >= 2:
        result = (
            SearchService.search_one(db, keyword, source, limit, offset, date_range=date_range, scrape_keyword=sk, group_url=gu)
            if source else
            SearchService.search_all(db, keyword, limit, offset, date_range=date_range, scrape_keyword=sk, group_url=gu)
        )
    else:
        result = (
            SearchService.recent_one(db, source, limit, offset, date_range=date_range, scrape_keyword=sk, group_url=gu)
            if source else
            SearchService.recent_all(db, limit, offset, date_range=date_range, scrape_keyword=sk, group_url=gu)
        )
    return {
        "status":  "success",
        "keyword": keyword or "",
        "source":  source or "all",
        "mode":    "search" if keyword else "recent",
        "limit":   limit,
        "offset":  offset,
        **result,
    }

@app.post("/export/selected", tags=["Search"])
async def export_selected(request: Request, db: Session = Depends(get_db)):
    from services.search_service import SearchService, ALL_SOURCES
    try:
        body = await request.json()
    except Exception:
        raise HTTPException(400, "Invalid JSON body")
    selections = body.get("selections", {})
    if not selections:
        raise HTTPException(400, "Provide 'selections': { source: [id, ...] }")
    bad = [s for s in selections if s not in ALL_SOURCES]
    if bad:
        raise HTTPException(400, f"Unknown source(s): {bad}")
    result = SearchService.export_selected(db, selections)
    return {
        "status":        "success",
        "total":         result["total"],
        "source_totals": result["source_totals"],
        "by_source":     result["by_source"],
    }

@app.get("/api/record/{source}/{record_id:path}", tags=["Search"])
def get_record(
    source:    str = FPath(...),
    record_id: str = FPath(...),
    db:        Session = Depends(get_db),
):
    """Fetch the full data for a single scraped record by source and native ID."""
    from services.search_service import _SELECTED_EXPORTERS, ALL_SOURCES
    if source not in ALL_SOURCES:
        raise HTTPException(400, f"Unknown source '{source}'. Valid: {ALL_SOURCES}")
    fn = _SELECTED_EXPORTERS.get(source)
    if fn is None:
        raise HTTPException(400, f"No exporter for source '{source}'")
    try:
        rows = fn(db, [record_id])
    except Exception as exc:
        raise HTTPException(500, str(exc))
    if not rows:
        raise HTTPException(404, "Record not found")
    return {"status": "success", "source": source, "data": rows[0]}

@app.delete("/api/record/{source}/{record_id:path}", tags=["Search"])
def delete_record(
    source:    str = FPath(...),
    record_id: str = FPath(...),
    db:        Session = Depends(get_db),
):
    """Delete a single scraped record by source and native ID."""
    from services.search_service import ALL_SOURCES
    if source not in ALL_SOURCES:
        raise HTTPException(400, f"Unknown source '{source}'. Valid: {ALL_SOURCES}")
    
    try:
        from services.search_service import _get_table_and_id_column
        from sqlalchemy import Integer as _SAInt
        table, id_col = _get_table_and_id_column(source)

        # Cast string path param to int when the DB column is integer-typed
        typed_id: object = record_id
        if isinstance(id_col.type, _SAInt):
            try:
                typed_id = int(record_id)
            except ValueError:
                raise HTTPException(400, f"record_id '{record_id}' must be an integer for source '{source}'")

        # Use ORM-session delete so cascade="all, delete-orphan" fires on child rows
        instance = db.query(table).filter(id_col == typed_id).first()
        if instance is None:
            raise HTTPException(404, "Record not found")
        db.delete(instance)
        db.commit()

        return {"status": "success", "message": "Deleted 1 record"}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(500, str(exc))


@app.get("/export", tags=["Search"])
def export_data(
    keyword: str | None = Query(None),
    source:  str | None = Query(None),
    db:      Session    = Depends(get_db),
):
    from services.search_service import SearchService, ALL_SOURCES
    if source and source not in ALL_SOURCES:
        raise HTTPException(400, f"Unknown source '{source}'. Valid: {ALL_SOURCES}")
    kw = keyword.strip() if keyword and keyword.strip() else None
    if source:
        result = SearchService.export_one(db, source, keyword=kw, limit=100_000)
        return {
            "status": "success", "keyword": kw or "",
            "source": source, "total": result["total"], "results": result["results"],
        }
    result = SearchService.export_all(db, keyword=kw, limit=100_000)
    return {
        "status": "success", "keyword": kw or "", "source": "all",
        "total": result["total"], "source_totals": result["source_totals"],
        "by_source": result["by_source"],
    }


# ══════════════════════════════════════════════════════════════════════════════
#  Results (JSON file browser)
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/api/results", tags=["Results"])
def list_results():
    files = sorted(
        [f for f in RESULTS_DIR.glob("*.json") if not f.name.startswith("seen_ids")],
        key=lambda f: f.stat().st_mtime, reverse=True,
    )
    return {
        "total": len(files),
        "files": [
            {
                "name":        f.name,
                "scraper":     f.name.split("_")[0],
                "size_kb":     round(f.stat().st_size / 1024, 1),
                "modified_at": datetime.fromtimestamp(
                    f.stat().st_mtime, tz=timezone.utc
                ).isoformat(),
            }
            for f in files
        ],
    }

@app.get("/api/results/download/{filename}", tags=["Results"])
def download_result(filename: str = FPath(...)):
    if ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(400, "Invalid filename.")
    path = RESULTS_DIR / filename
    if not path.exists():
        raise HTTPException(404, f"File '{filename}' not found.")
    return FileResponse(
        path=path, media_type="application/json", filename=filename,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

@app.get("/api/results/view/{filename}", tags=["Results"])
def view_result(filename: str = FPath(...)):
    if ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(400, "Invalid filename.")
    path = RESULTS_DIR / filename
    if not path.exists():
        raise HTTPException(404, f"File '{filename}' not found.")
    return JSONResponse(content=json.loads(path.read_text(encoding="utf-8")))

@app.get("/api/results/{scraper}", tags=["Results"])
def list_results_for_scraper(scraper: str = FPath(...)):
    if scraper not in VALID_SCRAPERS:
        raise HTTPException(400, f"Unknown scraper '{scraper}'. Valid: {sorted(VALID_SCRAPERS)}")
    files = sorted(
        RESULTS_DIR.glob(f"{scraper}_*.json"),
        key=lambda f: f.stat().st_mtime, reverse=True,
    )
    return {
        "scraper": scraper,
        "total":   len(files),
        "files": [
            {
                "name":        f.name,
                "size_kb":     round(f.stat().st_size / 1024, 1),
                "modified_at": datetime.fromtimestamp(
                    f.stat().st_mtime, tz=timezone.utc
                ).isoformat(),
            }
            for f in files
        ],
    }